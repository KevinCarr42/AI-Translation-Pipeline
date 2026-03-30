import json
import sys
import unicodedata
from copy import deepcopy
from datetime import datetime, timezone

import docx
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W = f'{{{W_NS}}}'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

_SMART_QUOTES = {
    '\u201c': '"', '\u201d': '"',
    '\u2018': "'", '\u2019': "'",
}


def _normalize_with_map(text):
    result = []
    omap = []
    prev_was_space = False
    for i, ch in enumerate(text):
        ch = _SMART_QUOTES.get(ch, ch)
        if unicodedata.category(ch) == 'Zs' or ch in ('\u00a0', '\ufeff'):
            ch = ' '
        if ch == ' ' and prev_was_space:
            continue
        omap.append(i)
        result.append(ch)
        prev_was_space = (ch == ' ')
    omap.append(len(text))
    return ''.join(result), omap


def _normalize(text):
    normalized, _ = _normalize_with_map(text)
    return normalized


def get_paragraph_text(para_elem):
    texts = []
    for t_elem in para_elem.iter(f'{W}t'):
        # Skip text inside existing deletions
        parent = t_elem.getparent()
        in_del = False
        node = parent
        while node is not None:
            if node.tag == f'{W}del':
                in_del = True
                break
            node = node.getparent()
        if not in_del:
            texts.append(t_elem.text or '')
    return ''.join(texts)


def get_direct_runs(para_elem):
    runs = []
    for child in para_elem:
        if child.tag == f'{W}r':
            # Skip runs inside del elements
            in_del = False
            node = child.getparent()
            while node is not None:
                if node.tag == f'{W}del':
                    in_del = True
                    break
                node = node.getparent()
            if not in_del:
                text = ''.join(t.text or '' for t in child.findall(f'{W}t'))
                runs.append((child, text))
        elif child.tag == f'{W}hyperlink':
            for run in child.findall(f'{W}r'):
                text = ''.join(t.text or '' for t in run.findall(f'{W}t'))
                runs.append((run, text))
    return runs


def make_run(text, rpr_template=None):
    run = etree.Element(f'{W}r')
    if rpr_template is not None:
        run.append(deepcopy(rpr_template))
    t = etree.SubElement(run, f'{W}t')
    t.set(XML_SPACE, 'preserve')
    t.text = text
    return run


def apply_tracked_change(para_elem, error_text, suggested_fix, change_id,
                         author='AI Review', date=None):
    if not error_text or not suggested_fix or error_text == suggested_fix:
        return False
    
    runs = get_direct_runs(para_elem)
    if not runs:
        return False
    
    full_text = ''.join(t for _, t in runs)
    norm_full, omap = _normalize_with_map(full_text)
    norm_error = _normalize(error_text)
    norm_idx = norm_full.find(norm_error)
    if norm_idx == -1:
        return False

    # Warn if error_text matches multiple times — we take the first occurrence
    if norm_full.find(norm_error, norm_idx + 1) != -1:
        snippet = error_text[:40] + '...' if len(error_text) > 40 else error_text
        print(f'  WARN: Multiple matches for "{snippet}" — using first occurrence')

    idx = omap[norm_idx]
    end_idx = omap[norm_idx + len(norm_error)]

    # Map character positions to runs
    pos = 0
    affected = []
    for run_elem, run_text in runs:
        run_start = pos
        run_end = pos + len(run_text)
        if run_end > idx and run_start < end_idx:
            clip_start = max(idx, run_start) - run_start
            clip_end = min(end_idx, run_end) - run_start
            affected.append((run_elem, run_text, clip_start, clip_end))
        pos = run_end
    
    if not affected:
        return False
    
    first_run, first_text, first_cs, first_ce = affected[0]
    last_run, last_text, last_cs, last_ce = affected[-1]
    
    # Get formatting from first affected run
    rpr = first_run.find(f'{W}rPr')
    
    # Determine insertion point
    parent = first_run.getparent()
    insert_idx = list(parent).index(first_run)
    
    new_elements = []
    
    # Text before the error in the first affected run
    if first_cs > 0:
        new_elements.append(make_run(first_text[:first_cs], rpr))
    
    # <w:del> element
    del_elem = etree.Element(f'{W}del')
    del_elem.set(f'{W}id', str(change_id))
    del_elem.set(f'{W}author', author)
    del_elem.set(f'{W}date', date)
    del_run = etree.SubElement(del_elem, f'{W}r')
    if rpr is not None:
        del_run.append(deepcopy(rpr))
    del_t = etree.SubElement(del_run, f'{W}delText')
    del_t.set(XML_SPACE, 'preserve')
    del_t.text = full_text[idx:end_idx]
    new_elements.append(del_elem)
    
    # <w:ins> element
    ins_elem = etree.Element(f'{W}ins')
    ins_elem.set(f'{W}id', str(change_id + 1))
    ins_elem.set(f'{W}author', author)
    ins_elem.set(f'{W}date', date)
    ins_run = etree.SubElement(ins_elem, f'{W}r')
    if rpr is not None:
        ins_run.append(deepcopy(rpr))
    ins_t = etree.SubElement(ins_run, f'{W}t')
    ins_t.set(XML_SPACE, 'preserve')
    ins_t.text = suggested_fix
    new_elements.append(ins_elem)
    
    # Text after the error in the last affected run
    if last_ce < len(last_text):
        new_elements.append(make_run(last_text[last_ce:], rpr))
    
    # Insert new elements before the first affected run
    for i, elem in enumerate(new_elements):
        parent.insert(insert_idx + i, elem)
    
    # Remove all affected runs
    for run_elem, _, _, _ in affected:
        run_parent = run_elem.getparent()
        run_parent.remove(run_elem)
    
    return True


def build_location_maps(doc):
    from scitrans.proofreader.extract_text import _iter_header_footer_paragraphs

    para_map = {}
    for i, p in enumerate(doc.paragraphs):
        para_map[f'P{i}'] = p._element

    table_map = {}
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            key = f'T{t_idx}-R{r_idx}'
            if key not in table_map:
                table_map[key] = []
            seen = set()
            for cell in row.cells:
                for p in cell.paragraphs:
                    elem_id = id(p._element)
                    if elem_id not in seen:
                        seen.add(elem_id)
                        table_map[key].append(p._element)

    hf_map = {}
    for loc_id, para in _iter_header_footer_paragraphs(doc):
        hf_map[loc_id] = para._element

    return para_map, table_map, hf_map


def main(translated_path, review_path, output_path):
    with open(review_path, 'r', encoding='utf-8') as f:
        errors = json.load(f)

    doc = docx.Document(translated_path)
    para_map, table_map, hf_map = build_location_maps(doc)

    change_id = 100
    applied = 0
    skipped = 0
    skipped_items = []
    date = datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')

    for error in errors:
        location = error['location']
        error_text = error.get('error_text', '')
        suggested_fix = error.get('suggested_fix', '')

        if not error_text or not suggested_fix or error_text == suggested_fix:
            skipped += 1
            skipped_items.append({**error, 'skip_reason': 'empty or no-op'})
            continue

        para_elems = []
        if location in para_map:
            para_elems = [para_map[location]]
        elif location in table_map:
            para_elems = table_map[location]
        elif location in hf_map:
            para_elems = [hf_map[location]]
        else:
            print(f'  SKIP: Location {location} not found in document')
            skipped += 1
            skipped_items.append({**error, 'skip_reason': 'location not found'})
            continue

        success = False
        for para_elem in para_elems:
            if apply_tracked_change(para_elem, error_text, suggested_fix,
                                    change_id, date=date):
                success = True
                change_id += 2
                break

        if success:
            applied += 1
        else:
            snippet = error_text[:60] + '...' if len(error_text) > 60 else error_text
            print(f'  SKIP: Could not match text in {location}: "{snippet}"')
            skipped += 1
            skipped_items.append({**error, 'skip_reason': 'text not found'})

    doc.save(output_path)
    print(f'\nDone: {applied} changes applied, {skipped} skipped')
    print(f'Saved to: {output_path}')
    return {'applied': applied, 'skipped': skipped, 'skipped_items': skipped_items}


if __name__ == '__main__':
    if len(sys.argv) != 4:
        print('Usage: python apply_review.py <translated.docx> <review.json> <output.docx>')
        print('Example: python apply_review.py to_review/1432_en_translated.docx to_review/1432_en_review.json to_review/1432_en_recommended_updates.docx')
        sys.exit(1)
    main(sys.argv[1], sys.argv[2], sys.argv[3])
