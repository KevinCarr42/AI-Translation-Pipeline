import docx


def _iter_header_footer_paragraphs(document):
    seen_ids = set()
    header_idx = 0
    footer_idx = 0
    
    header_footer_attrs = [
        ('header', 'H'),
        ('first_page_header', 'H'),
        ('even_page_header', 'H'),
        ('footer', 'F'),
        ('first_page_footer', 'F'),
        ('even_page_footer', 'F'),
    ]
    
    for section in document.sections:
        for attr, prefix in header_footer_attrs:
            hf = getattr(section, attr)
            if id(hf._element) in seen_ids or hf.is_linked_to_previous:
                continue
            seen_ids.add(id(hf._element))
            
            for paragraph in hf.paragraphs:
                if not paragraph.text.strip():
                    continue
                if prefix == 'H':
                    yield f'H{header_idx}', paragraph.text
                    header_idx += 1
                else:
                    yield f'F{footer_idx}', paragraph.text
                    footer_idx += 1


def extract_text_with_ids(filepath):
    doc = docx.Document(filepath)
    lines = []
    
    # Paragraphs
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            lines.append(f'[P{i}] {p.text}')
    
    # Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells = [c.text for c in row.cells]
            lines.append(f'[T{t_idx}-R{r_idx}] ' + ' | '.join(cells))
    
    # Headers and footers
    for loc_id, text in _iter_header_footer_paragraphs(doc):
        lines.append(f'[{loc_id}] {text}')
    
    return '\n'.join(lines)


def extract_locations(filepath):
    doc = docx.Document(filepath)
    locations = []
    
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            locations.append((f'P{i}', p.text))
    
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells = [c.text for c in row.cells]
            locations.append((f'T{t_idx}-R{r_idx}', ' | '.join(cells)))
    
    for loc_id, text in _iter_header_footer_paragraphs(doc):
        locations.append((loc_id, text))
    
    return locations
