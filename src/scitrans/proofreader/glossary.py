import json
import re

import docx

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W = f'{{{W_NS}}}'


def detect_language(doc):
    body = doc.element.body
    for rpr in body.iter(f'{W}rPr'):
        lang_elem = rpr.find(f'{W}lang')
        if lang_elem is not None:
            val = lang_elem.get(f'{W}val', '')
            if val.startswith('fr'):
                return 'fr'
            if val.startswith('en'):
                return 'en'
    
    for style in doc.styles:
        if style.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH:
            elem = style.element
            rpr = elem.find(f'{W}rPr')
            if rpr is not None:
                lang_elem = rpr.find(f'{W}lang')
                if lang_elem is not None:
                    val = lang_elem.get(f'{W}val', '')
                    if val.startswith('fr'):
                        return 'fr'
                    if val.startswith('en'):
                        return 'en'
    return None


def detect_language_from_path(filepath):
    doc = docx.Document(filepath)
    return detect_language(doc)


def load_glossary(filepath, categories=None, source_lang=None):
    with open(filepath, 'r', encoding='utf-8') as f:
        data = json.load(f)
    all_categories = ["nomenclature", "taxon", "site", "table", "acronym"]
    if categories is None:
        categories = all_categories
    
    if source_lang not in ("en", "fr"):
        raise ValueError(f'source_lang must be "en" or "fr", got {source_lang!r}')
    
    if source_lang == "en":
        src_key, tgt_key = "english", "french"
        src_acronym_key, tgt_acronym_key = "english_acronym", "french_acronym"
    else:
        src_key, tgt_key = "french", "english"
        src_acronym_key, tgt_acronym_key = "french_acronym", "english_acronym"
    
    glossary = {}
    for category in categories:
        if category == "acronym":
            for entry in data.get("translations", {}).get("acronym", []):
                src = entry.get(src_acronym_key) or entry.get(src_key)
                tgt = entry.get(tgt_acronym_key) or entry.get(tgt_key)
                if src and tgt:
                    glossary[src] = tgt
        else:
            for entry in data.get("translations", {}).get(category, []):
                if entry.get(src_key) and entry.get(tgt_key):
                    glossary[entry[src_key]] = entry[tgt_key]
    return glossary


def extract_text(filepath):
    return "\n".join([p.text for p in docx.Document(filepath).paragraphs])


def build_sub_glossary(text, glossary):
    text_lower = text.lower()
    sub_glossary = {}
    for en, fr in glossary.items():
        if re.search(rf'\b{re.escape(str(en).lower())}\b', text_lower):
            sub_glossary[en] = fr
    return sub_glossary
