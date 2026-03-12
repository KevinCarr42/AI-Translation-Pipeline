import json
import os
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.ns import qn as oxml_qn
from docx.shared import Inches

from scitrans.translate.word_formatting import FormattedRun, RuleRegistry


def strip_hyperlink_xml(paragraph):
    p_elem = paragraph._element
    hyperlink_elems = list(p_elem.findall(qn('w:hyperlink')))
    if not hyperlink_elems:
        return []
    
    hyperlink_data = []
    for hl_elem in hyperlink_elems:
        r_id = hl_elem.get(qn('r:id'))
        url = paragraph.part.rels[r_id].target_ref if r_id and r_id in paragraph.part.rels else ''
        for r_elem in hl_elem.findall(qn('w:r')):
            t_elem = r_elem.find(qn('w:t'))
            text = t_elem.text if t_elem is not None and t_elem.text else ''
            hyperlink_data.append((text, url))
    
    for hl_elem in hyperlink_elems:
        for r_elem in list(hl_elem.findall(qn('w:r'))):
            p_elem.insert(list(p_elem).index(hl_elem), r_elem)
        p_elem.remove(hl_elem)
    
    return hyperlink_data


def extract_hyperlink_notes(paragraph, formatting_records, location=None):
    hyperlink_data = strip_hyperlink_xml(paragraph)
    if not hyperlink_data:
        return False
    
    full_paragraph_text = paragraph.text
    for original_text, url in hyperlink_data:
        record = {
            'original_text': original_text,
            'full_paragraph': full_paragraph_text,
            'notes': url,
            'type': 'url',
        }
        if location:
            record['location'] = location
        formatting_records.append(record)
    
    return True


def _filter_notes(records):
    filtered = []
    for record in records:
        original = record.get('original_text', '')
        if not original or not original.strip():
            continue
        notes = record.get('notes', '')
        if not notes or not notes.strip():
            continue
        note_lines = [line.strip() for line in notes.split('\n') if line.strip()]
        meaningful = [
            line for line in note_lines
            if line != 'colour=000000' and line.strip() not in ('', ' ')
        ]
        if not meaningful:
            continue
        record = dict(record)
        record['notes'] = '\n'.join(meaningful)
        filtered.append(record)
    return filtered


def _location_section(record):
    location = record.get('location')
    if not location:
        return None
    return (location.get('section'), location.get('index'), location.get('table'), location.get('row'), location.get('cell'))


def _group_notes_by_paragraph(records):
    groups = {}
    for record in records:
        key = record.get('full_paragraph', '')
        if key not in groups:
            groups[key] = []
        dedup_key = (record.get('notes', ''), _location_section(record))
        existing = [(r.get('notes', ''), _location_section(r)) for r in groups[key]]
        if dedup_key not in existing:
            groups[key].append(record)
    return groups


def _get_note_type(details_list):
    has_formatting = any(d.get('type') == 'formatting' for d in details_list)
    has_url = any(d.get('type') == 'url' for d in details_list)
    if has_formatting and has_url:
        return 'mixed'
    if has_url:
        return 'url'
    return 'formatting'


def write_notes_json(formatting_records, output_path):
    filtered = _filter_notes(formatting_records)
    if not filtered:
        return
    
    grouped = _group_notes_by_paragraph(filtered)
    paragraphs_section = []
    tables_section = []
    
    for full_paragraph, details in grouped.items():
        location = next((d['location'] for d in details if 'location' in d), None)
        entry = {
            'full_paragraph': full_paragraph,
            'type': _get_note_type(details),
            'notes': [
                {
                    'original_text': d.get('original_text', ''),
                    'detail': d.get('notes', ''),
                    'type': d.get('type', 'formatting'),
                }
                for d in details
            ],
        }
        if location:
            entry['location'] = location
        
        if location and location.get('section') == 'tables':
            tables_section.append(entry)
        else:
            paragraphs_section.append(entry)
    
    output = {
        'paragraphs': paragraphs_section,
        'tables': tables_section,
    }
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(output, f, indent=2, ensure_ascii=False)


def add_formatting_notes(paragraph, formatting_records, detected_rules=None, location=None):
    full_paragraph_text = paragraph.text
    
    for run in list(paragraph.runs):
        formatted_run = FormattedRun.create(run)
        
        if formatted_run.has_formatting and formatted_run.text.strip():
            if detected_rules and RuleRegistry.is_auto_handled(formatted_run, detected_rules):
                continue
            record = {
                'original_text': run.text,
                'full_paragraph': full_paragraph_text,
                'notes': formatted_run.formatting_notes,
                'type': 'formatting',
            }
            if location:
                record['location'] = location
            formatting_records.append(record)


def json_to_word_tables(json_file, delete_json=False):
    json_path = Path(json_file)
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    COLOR_YELLOW = "FFFF00"
    COLOR_CYAN = "00FFFF"
    COLOR_GREEN = "00FF00"
    
    def _cell_color_for_notes(notes):
        types = {n.get('type', 'formatting') for n in notes}
        if 'url' in types and 'formatting' in types:
            return COLOR_GREEN
        if 'url' in types:
            return COLOR_CYAN
        return COLOR_YELLOW
    
    def _shade_cell(cell, hex_color):
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(oxml_qn('w:shd'), {
            oxml_qn('w:fill'): hex_color,
            oxml_qn('w:val'): 'clear',
        })
        shading.append(shd)
    
    def _build_table(entries, heading_text):
        doc.add_heading(heading_text, level=1)
        if not entries:
            doc.add_paragraph("No notes.")
            return
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        header_row = table.rows[0]
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        header_row.cells[0].paragraphs[0].add_run("Full Paragraph (original language)").bold = True
        header_row.cells[1].paragraphs[0].add_run("Details").bold = True
        
        row_elem = header_row._tr
        tr_pr = row_elem.get_or_add_trPr()
        tr_pr.append(tr_pr.makeelement(oxml_qn('w:tblHeader'), {}))
        
        for entry in entries:
            row = table.add_row()
            row.cells[0].text = entry.get('full_paragraph', '')
            
            notes = entry.get('notes', [])
            details_cell = row.cells[1]
            details_cell.text = ''
            for i, note in enumerate(notes):
                original = note.get('original_text', '')
                detail = note.get('detail', '')
                text = f'"{original}": {detail}'
                p = details_cell.paragraphs[0] if i == 0 else details_cell.add_paragraph()
                p.style = 'List Bullet'
                p.text = text
            
            if notes:
                _shade_cell(details_cell, _cell_color_for_notes(notes))
        
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            tr_pr.append(tr_pr.makeelement(oxml_qn('w:cantSplit'), {}))
    
    _build_table(data.get('paragraphs', []), 'Paragraphs')
    _build_table(data.get('tables', []), 'Tables')
    
    output_path = json_path.with_suffix('.docx')
    doc.save(str(output_path))
    
    if delete_json:
        os.remove(json_path)
