import copy
import json
import os
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scitrans import config
from scitrans.translate.models import create_translator
from scitrans.translate.utils import normalize_apostrophes
from scitrans.translate.utils import split_into_chunks, reassemble_sentences, reassemble_paragraphs
from scitrans.translate.word_formatting import FormattedRun
from scitrans.translate.word_formatting import apply_formatting_rules, RuleRegistry
from scitrans.translate.word_notes import add_formatting_notes, extract_hyperlink_notes, write_notes_json, json_to_word_tables
from scitrans.translate.word_formatting import is_numeric, convert_numeric, parse_formatted_string
from scitrans.rules_based_replacements.token_utils import get_translation_value, normalize_translations

_MC_NS = 'http://schemas.openxmlformats.org/markup-compatibility/2006'


def _iter_document_elements(document):
    for para_idx, paragraph in enumerate(document.paragraphs):
        yield paragraph, {"section": "paragraphs", "index": para_idx}, "paragraph"
    
    for table_idx, table in enumerate(document.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                yield cell, {"section": "tables", "table": table_idx, "row": row_idx, "cell": cell_idx}, "cell"
    
    translated_hf_ids = set()
    header_footer_attrs = [
        'header', 'footer', 'first_page_header', 'first_page_footer',
        'even_page_header', 'even_page_footer',
    ]
    
    for section in document.sections:
        for attr in header_footer_attrs:
            hf = getattr(section, attr)
            if id(hf._element) in translated_hf_ids or hf.is_linked_to_previous:
                continue
            
            translated_hf_ids.add(id(hf._element))
            
            for paragraph in hf.paragraphs:
                yield paragraph, {"section": "headers_footers", "type": attr}, "paragraph"
            
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield cell, {"section": "headers_footers", "type": attr, "in_table": True}, "cell"


def _has_formatting_differences(paragraph):
    if not paragraph.runs:
        return False
    first_format = FormattedRun.create(paragraph.runs[0])
    return any(run.text.strip() and FormattedRun.create(run) != first_format for run in paragraph.runs[1:])


def _isolate_run_tabs(paragraph):
    for run in list(paragraph.runs):
        if '\t' in run.text and run.text != '\t':
            parts = run.text.split('\t')
            for i, part in enumerate(parts):
                if part:
                    new_r = copy.deepcopy(run._element)
                    for child in list(new_r):
                        if child.tag in (qn('w:t'), qn('w:tab')):
                            new_r.remove(child)
                    t = OxmlElement('w:t')
                    t.text = part
                    if part.startswith(' ') or part.endswith(' '):
                        t.set(qn('xml:space'), 'preserve')
                    new_r.append(t)
                    run._element.addprevious(new_r)
                if i < len(parts) - 1:
                    tab_r = copy.deepcopy(run._element)
                    for child in list(tab_r):
                        if child.tag in (qn('w:t'), qn('w:tab')):
                            tab_r.remove(child)
                    tab_r.append(OxmlElement('w:tab'))
                    run._element.addprevious(tab_r)
            paragraph._element.remove(run._element)


def _collapse_runs_preserving_shapes(paragraph):
    if len(paragraph.runs) < 2:
        return
    p_element = paragraph._element
    runs = list(paragraph.runs)
    current_run = runs[0]
    
    for next_run in runs[1:]:
        current_elem = current_run._element
        next_elem = next_run._element
        
        has_special = (
                next_elem.findall(qn('w:drawing')) or current_elem.findall(qn('w:drawing')) or
                next_elem.findall(qn('w:fldChar')) or current_elem.findall(qn('w:fldChar')) or
                next_elem.findall(qn('w:instrText')) or current_elem.findall(qn('w:instrText')) or
                next_elem.findall(qn('w:tab')) or current_elem.findall(qn('w:tab'))
        )
        if has_special:
            current_run = next_run
            continue
        
        if FormattedRun.create(current_run) == FormattedRun.create(next_run):
            current_run.text += next_run.text
            p_element.remove(next_run._element)
        else:
            current_run = next_run


def _remove_orphaned_field_runs(paragraph):
    p_elem = paragraph._element
    paired = set()
    runs_with_fld = []
    depth = 0
    for child in p_elem:
        if child.tag != qn('w:r'):
            continue
        fld_char = child.find(qn('w:fldChar'))
        if fld_char is None:
            continue
        fld_type = fld_char.get(qn('w:fldCharType'))
        runs_with_fld.append((child, fld_type))
        if fld_type == 'begin':
            depth += 1
            paired.add(id(child))
        elif fld_type == 'end' and depth > 0:
            depth -= 1
            paired.add(id(child))

    for run_elem, fld_type in runs_with_fld:
        if id(run_elem) in paired:
            continue
        t = run_elem.find(qn('w:t'))
        if t is not None and (t.text or '').strip():
            continue
        p_elem.remove(run_elem)


def _has_only_field_runs(paragraph):
    p_elem = paragraph._element
    in_field = 0
    has_any_run = False
    for child in p_elem:
        if child.tag != qn('w:r'):
            continue
        has_any_run = True
        fld_char = child.find(qn('w:fldChar'))
        if fld_char is not None:
            fld_type = fld_char.get(qn('w:fldCharType'))
            if fld_type == 'begin':
                in_field += 1
            elif fld_type == 'end':
                in_field -= 1
            continue
        if in_field > 0:
            continue
        t_elem = child.find(qn('w:t'))
        if t_elem is not None and (t_elem.text or '').strip():
            return False
    return has_any_run


def _unwrap_smart_tags(paragraph):
    for smart_tag in list(paragraph._element.findall(qn('w:smartTag'))):
        parent = smart_tag.getparent()
        for child in list(smart_tag):
            if child.tag == qn('w:r'):
                smart_tag.addprevious(child)
        parent.remove(smart_tag)


def _extract_non_run_elements(paragraph):
    _unwrap_smart_tags(paragraph)
    saved = []
    for child in list(paragraph._element):
        if child.tag == qn('w:pPr'):
            continue
        if child.tag != qn('w:r'):
            saved.append(child)
            paragraph._element.remove(child)
        elif child.findall(qn('w:drawing')) or child.findall(qn('m:oMath')) or child.findall(qn('w:object')):
            saved.append(child)
            paragraph._element.remove(child)
    return saved


def _chunk_and_translate(source_text, translation_manager, source_lang, target_lang, use_find_replace, idx, use_cache, preferential_dict, chunk_by):
    chunks, chunk_metadata = split_into_chunks(source_text, chunk_by)
    translated_chunks = []
    
    for i, chunk in enumerate(chunks, idx):
        if not chunk.strip():
            translated_chunks.append('')
            continue
        
        result = translation_manager.translate_with_best_model(
            text=chunk,
            source_lang=source_lang,
            target_lang=target_lang,
            use_find_replace=use_find_replace,
            idx=i,
            use_cache=use_cache
        )
        translated_chunks.append(result.get("translated_text", chunk))
    
    if chunk_by == "paragraphs":
        return reassemble_paragraphs(translated_chunks, chunk_metadata)
    return reassemble_sentences(translated_chunks, chunk_metadata)


def _reinsert_non_run_elements(paragraph, saved_elements):
    for elem in saved_elements:
        paragraph._element.append(elem)


def _convert_newlines_to_breaks(paragraph):
    for run in paragraph.runs:
        if '\n' in run.text:
            parts = run.text.split('\n')
            run.text = parts[0]
            
            for part in parts[1:]:
                br = OxmlElement('w:br')
                run._element.append(br)
                if part:
                    t = OxmlElement('w:t')
                    t.text = part
                    if part.startswith(' ') or part.endswith(' '):
                        t.set(qn('xml:space'), 'preserve')
                    run._element.append(t)


def _group_text_runs_between_tabs(paragraph):
    groups = []
    current_group = []
    for run in paragraph.runs:
        if run._element.findall(qn('w:tab')):
            if current_group:
                groups.append(current_group)
                current_group = []
            groups.append(None)
        else:
            current_group.append(run)
    if current_group:
        groups.append(current_group)
    return groups


def _translate_paragraph(
        paragraph, translation_manager, source_lang, target_lang,
        use_find_replace, idx, use_cache=True, formatting_records=None,
        preferential_dict=None, chunk_by="sentences", location=None
):
    has_hl = extract_hyperlink_notes(paragraph, formatting_records, location=location)
    detected = RuleRegistry.detect_all(paragraph)
    has_fmt = _has_formatting_differences(paragraph)
    records_before = len(formatting_records) if formatting_records is not None else 0
    
    if has_fmt:
        add_formatting_notes(paragraph, formatting_records, detected_rules=detected, location=location)
    
    source_text = paragraph.text
    if not source_text or not source_text.strip():
        return idx
    
    saved_elements = _extract_non_run_elements(paragraph)
    _remove_orphaned_field_runs(paragraph)

    _isolate_run_tabs(paragraph)
    _collapse_runs_preserving_shapes(paragraph)
    
    if _has_only_field_runs(paragraph):
        _reinsert_non_run_elements(paragraph, saved_elements)
        return idx
    
    groups = _group_text_runs_between_tabs(paragraph)
    has_tabs = any(g is None for g in groups)
    
    if has_tabs:
        for group in groups:
            if group is None:
                continue
            group_text = ''.join(run.text or '' for run in group)
            if not group_text.strip():
                continue
            translated_segment = _chunk_and_translate(
                group_text, translation_manager, source_lang, target_lang,
                use_find_replace, idx, use_cache, preferential_dict=preferential_dict,
                chunk_by=chunk_by
            )
            group[0].text = normalize_apostrophes(translated_segment)
            for run in group[1:]:
                run.text = ''
    else:
        full_text = paragraph.text
        if full_text and full_text.strip():
            translated_text = _chunk_and_translate(
                full_text, translation_manager, source_lang, target_lang,
                use_find_replace, idx, use_cache, preferential_dict=preferential_dict,
                chunk_by=chunk_by
            )
            normalized = normalize_apostrophes(translated_text)
            runs = paragraph.runs
            if runs:
                runs[0].text = normalized
                for run in runs[1:]:
                    run.text = ''
            else:
                paragraph.text = normalized
    
    _reinsert_non_run_elements(paragraph, saved_elements)
    
    apply_formatting_rules(paragraph, formatting_records, source_text, location=location, detected=detected)
    
    has_fmt_notes = (len(formatting_records) if formatting_records is not None else 0) > records_before
    
    if has_hl or has_fmt_notes:
        color = WD_COLOR_INDEX.TURQUOISE
        if has_fmt_notes and not has_hl:
            color = WD_COLOR_INDEX.YELLOW
        elif has_fmt_notes and has_hl:
            color = WD_COLOR_INDEX.BRIGHT_GREEN
        for run in paragraph.runs:
            if run.text and run.text.strip():
                run.font.highlight_color = color
    
    _convert_newlines_to_breaks(paragraph)
    return idx + 1


def _translate_table_cell(
        cell, translation_manager, source_lang, target_lang,
        use_find_replace, idx, use_cache=True, formatting_records=None,
        preferential_dict=None, table_translations_dict=None,
        chunk_by="sentences", location=None
):
    to_fr = target_lang == "fr"
    cell_text = cell.text
    
    if not cell_text or not cell_text.strip():
        return idx
    
    stripped = cell_text.strip()
    
    if config.NUMERIC_CONVERSION_CONFIG.get("enabled") and is_numeric(stripped):
        converted = convert_numeric(stripped, to_fr=to_fr)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.text.strip():
                    run.text = converted
                    converted = ""
    elif table_translations_dict and stripped in table_translations_dict:
        raw_replacement = table_translations_dict[stripped]
        formatted_runs = parse_formatted_string(raw_replacement)
        content_runs = [run for p in cell.paragraphs for run in p.runs if run.text.strip()]
        for run in cell.paragraphs[0].runs:
            run.text = ''
        for i, fmt_run in enumerate(formatted_runs):
            if i < len(content_runs):
                content_runs[i].text = fmt_run.text
                content_runs[i].italic = fmt_run.italic
                if fmt_run.superscript:
                    content_runs[i].font.superscript = True
                if fmt_run.subscript:
                    content_runs[i].font.subscript = True
            else:
                content_runs[-1].text += fmt_run.text
    elif preferential_dict:
        match_translation = _find_preferential_match(stripped, source_lang, preferential_dict)
        if match_translation:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.text.strip():
                        run.text = match_translation
                        match_translation = ""
        elif len(stripped) >= config.TABLE_TRANSLATION_CONFIG.get("min_cell_length_for_ai", 20):
            for paragraph in cell.paragraphs:
                idx = _translate_paragraph(
                    paragraph, translation_manager, source_lang, target_lang,
                    use_find_replace, idx, use_cache=use_cache,
                    formatting_records=formatting_records,
                    preferential_dict=preferential_dict,
                    chunk_by=chunk_by, location=location
                )
            return idx
    elif len(stripped) >= config.TABLE_TRANSLATION_CONFIG.get("min_cell_length_for_ai", 20):
        for paragraph in cell.paragraphs:
            idx = _translate_paragraph(
                paragraph, translation_manager, source_lang, target_lang,
                use_find_replace, idx, use_cache=use_cache,
                formatting_records=formatting_records,
                preferential_dict=preferential_dict,
                chunk_by=chunk_by, location=location
            )
        return idx
    
    return idx


def _find_preferential_match(stripped, source_lang, preferential_dict):
    lookup_key = stripped.lower()
    pref_translations = preferential_dict.get("translations", preferential_dict)
    pref_translations = normalize_translations(pref_translations)
    for category, terms in pref_translations.items():
        for term_key, term_data in terms.items():
            if term_key.lower() == lookup_key:
                if source_lang == "en":
                    match_translation = get_translation_value(term_data)
                else:
                    match_translation = term_key if source_lang == "fr" else None
                if match_translation:
                    if stripped[0].isupper() and match_translation[0].islower():
                        match_translation = match_translation[0].upper() + match_translation[1:]
                    return match_translation
    return None


def _set_proofing_language(document, target_lang):
    locale_map = {'fr': 'fr-CA', 'en': 'en-CA'}
    locale_code = locale_map[target_lang]
    
    elements = [document.element]
    
    header_footer_attrs = [
        'header', 'footer',
        'first_page_header', 'first_page_footer',
        'even_page_header', 'even_page_footer',
    ]
    seen_ids = set()
    for section in document.sections:
        for attr in header_footer_attrs:
            hf = getattr(section, attr)
            if id(hf._element) in seen_ids or hf.is_linked_to_previous:
                continue
            seen_ids.add(id(hf._element))
            elements.append(hf._element)
    
    for root_elem in elements:
        for r_elem in root_elem.iter(qn('w:r')):
            rPr = r_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r_elem.insert(0, rPr)
            lang = rPr.find(qn('w:lang'))
            if lang is None:
                lang = OxmlElement('w:lang')
                rPr.append(lang)
            lang.set(qn('w:val'), locale_code)


def translate_word_document(
        input_docx_file, output_docx_file=None, source_lang="en", chunk_by="sentences",
        models_to_use=None, use_find_replace=True, use_finetuned=True,
        translation_manager=None, include_timestamp=True, use_cache=True,
        preserve_json_notes=False
):
    if not output_docx_file:
        base, ext = os.path.splitext(input_docx_file)
        date_str = datetime.now().strftime("%Y%m%d")
        output_docx_file = f"{base}_translated_{date_str}.docx" if include_timestamp else f"{base}_translated.docx"
    
    if source_lang not in ["en", "fr"]:
        raise ValueError('source_lang must be either "fr" or "en"')
    
    target_lang = "fr" if source_lang == "en" else "en"
    
    if not translation_manager:
        translation_manager = create_translator(
            use_finetuned=use_finetuned, models_to_use=models_to_use,
            use_embedder=True, load_models=True
        )
    
    document = Document(input_docx_file)
    idx = 1
    formatting_records = []
    
    preferential_dict = None
    if os.path.exists(config.PREFERENTIAL_JSON_PATH):
        with open(config.PREFERENTIAL_JSON_PATH, 'r', encoding='utf-8') as f:
            preferential_dict = json.load(f)
    
    table_translations_dict = None
    if preferential_dict:
        table_entries = preferential_dict.get("translations", {}).get("table", [])
        if table_entries:
            source_key = "english" if source_lang == "en" else "french"
            target_formatted_key = "fr_formatted" if target_lang == "fr" else "en_formatted"
            table_translations_dict = {}
            for entry in table_entries:
                plain_key = entry.get(source_key, "")
                formatted_value = entry.get(target_formatted_key, "")
                if plain_key and formatted_value and plain_key not in table_translations_dict:
                    table_translations_dict[plain_key] = formatted_value
    
    for element, location, elem_type in _iter_document_elements(document):
        if elem_type == "paragraph":
            idx = _translate_paragraph(
                element, translation_manager, source_lang, target_lang,
                use_find_replace, idx, use_cache, formatting_records,
                preferential_dict, chunk_by, location
            )
        elif elem_type == "cell":
            idx = _translate_table_cell(
                element, translation_manager, source_lang, target_lang,
                use_find_replace, idx, use_cache, formatting_records,
                preferential_dict, table_translations_dict, chunk_by, location
            )
    
    _set_proofing_language(document, target_lang)
    document.save(output_docx_file)
    
    if formatting_records:
        json_notes_path = os.path.splitext(output_docx_file)[0] + '_translation_notes.json'
        write_notes_json(formatting_records, json_notes_path)
        json_to_word_tables(json_notes_path, preserve_json_notes=preserve_json_notes)
    
    return output_docx_file
