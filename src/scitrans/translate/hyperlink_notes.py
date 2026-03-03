from docx import Document
from docx.shared import Pt


def write_hyperlink_notes(hyperlink_records, output_path):
    document = Document()
    document.add_heading('Hyperlink Translation Notes', level=1)
    
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Original Text'
    header_cells[1].text = 'Full Sentence'
    header_cells[2].text = 'URL'
    
    for record in hyperlink_records:
        row_cells = table.add_row().cells
        row_cells[0].text = record['original_text']
        row_cells[1].text = record['full_sentence']
        row_cells[2].text = record['url']
    
    document.save(output_path)
