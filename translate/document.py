import logging
import re
from translate.models import create_translator

logger = logging.getLogger(__name__)


def split_by_sentences(text):
    lines = text.split('\n')
    chunks = []
    chunk_metadata = []
    
    for line_idx, line in enumerate(lines):
        if not line.strip():
            chunks.append('')
            chunk_metadata.append({
                'line_idx': line_idx,
                'sent_idx': 0,
                'is_last_in_line': True,
                'is_empty': True
            })
            continue
        
        sentences = _split_into_sentences(line)
        for sent_idx, sentence in enumerate(sentences):
            sentence = sentence.strip()
            if sentence:
                chunks.append(sentence)
                chunk_metadata.append({
                    'line_idx': line_idx,
                    'sent_idx': sent_idx,
                    'is_last_in_line': sent_idx == len(sentences) - 1,
                    'is_empty': False
                })
    
    return chunks, chunk_metadata


def reassemble_sentences(translated_chunks, chunk_metadata):
    lines_dict = {}
    for i, (translated_chunk, metadata) in enumerate(zip(translated_chunks, chunk_metadata)):
        line_idx = metadata['line_idx']
        if line_idx not in lines_dict:
            lines_dict[line_idx] = []
        
        lines_dict[line_idx].append(translated_chunk)
    
    for line_idx in lines_dict:
        if isinstance(lines_dict[line_idx], list):
            lines_dict[line_idx] = ' '.join(lines_dict[line_idx])
    
    return '\n'.join(lines_dict[i] for i in sorted(lines_dict.keys()))


def split_by_paragraphs(text):
    # Notes:
    #  still get >512 token issues with 1000 characters, use 600 to be conservative
    MAX_CHAR = 600
    
    lines = text.split('\n')
    chunks = []
    chunk_metadata = []
    
    para_idx = 0
    for line_idx, line in enumerate(lines):
        if not line.strip():
            chunks.append('')
            chunk_metadata.append({
                'line_idx': line_idx,
                'para_idx': para_idx,
                'is_empty': True
            })
            para_idx += 1
            continue
        
        if len(line) <= MAX_CHAR:
            chunks.append(line)
            chunk_metadata.append({
                'line_idx': line_idx,
                'para_idx': para_idx,
                'is_last_in_line': True,
                'is_empty': False
            })
        else:
            sentences = _split_into_sentences(line)
            line_chunks = []
            current_chunk = ''
            
            for sentence in sentences:
                if len(current_chunk) + len(sentence) + 1 <= MAX_CHAR:
                    current_chunk += (' ' if current_chunk else '') + sentence
                else:
                    if current_chunk:
                        line_chunks.append(current_chunk)
                    current_chunk = sentence
            
            if current_chunk:
                line_chunks.append(current_chunk)
            
            for chunk_idx, chunk in enumerate(line_chunks):
                chunks.append(chunk)
                chunk_metadata.append({
                    'line_idx': line_idx,
                    'para_idx': para_idx,
                    'is_last_in_line': chunk_idx == len(line_chunks) - 1,
                    'is_empty': False
                })
    
    return chunks, chunk_metadata


def reassemble_paragraphs(translated_chunks, chunk_metadata):
    lines_dict = {}
    for translated_chunk, metadata in zip(translated_chunks, chunk_metadata):
        line_idx = metadata['line_idx']
        if line_idx not in lines_dict:
            lines_dict[line_idx] = []
        
        lines_dict[line_idx].append(translated_chunk)
    
    for line_idx in lines_dict:
        if isinstance(lines_dict[line_idx], list):
            lines_dict[line_idx] = ' '.join(lines_dict[line_idx])
    
    return '\n'.join(lines_dict[i] for i in sorted(lines_dict.keys()))


def normalize_apostrophes(text):
    return text.replace("'", "'").replace("'", "'")


_PROTECTED_LABEL_PATTERN = re.compile(
    r'\b(?:Figure|Fig|Table|Tableau)\.?\s*\d+\.?', re.IGNORECASE
)
_PLACEHOLDER = '\x00'


def _split_into_sentences(text):
    protected_positions = []
    for match in _PROTECTED_LABEL_PATTERN.finditer(text):
        for i, ch in enumerate(match.group()):
            if ch == '.':
                protected_positions.append(match.start() + i)
    
    chars = list(text)
    for pos in protected_positions:
        chars[pos] = _PLACEHOLDER
    
    sentences = re.split(r'(?<=[.!?])\s+', ''.join(chars))
    return [s.replace(_PLACEHOLDER, '.') for s in sentences]


def translate_txt_document(
        input_text_file,
        output_text_file=None,
        source_lang="en",
        chunk_by="sentences",
        models_to_use=None,
        use_find_replace=True,
        use_finetuned=True,
        translation_manager=None,
        start_idx=0,
        single_attempt=False,
        use_cache=True
):
    if not output_text_file:
        import os
        base, ext = os.path.splitext(input_text_file)
        output_text_file = f"{base}_translated{ext}"
    
    if source_lang not in ["en", "fr"]:
        raise ValueError('source_lang must be either "fr" or "en"')
    
    target_lang = "fr" if source_lang == "en" else "en"
    
    with open(input_text_file, 'r', encoding='utf-8') as f:
        text = f.read()
    
    if chunk_by == "sentences":
        chunks, chunk_metadata = split_by_sentences(text)
    else:
        chunks, chunk_metadata = split_by_paragraphs(text)
    
    if not translation_manager:
        translation_manager = create_translator(
            use_finetuned=use_finetuned,
            models_to_use=models_to_use,
            use_embedder=True,
            load_models=True
        )
    
    translated_chunks = []
    for i, (chunk, metadata) in enumerate(zip(chunks, chunk_metadata), start_idx + 1):
        if metadata.get('is_empty', False):
            translated_chunks.append('')
            continue
        
        result = translation_manager.translate_with_best_model(
            text=chunk,
            source_lang=source_lang,
            target_lang=target_lang,
            use_find_replace=use_find_replace,
            idx=i,
            single_attempt=single_attempt,
            use_cache=use_cache
        )
        
        translated_text = result.get("translated_text", "[TRANSLATION FAILED]")
        translated_text = normalize_apostrophes(translated_text)
        translated_chunks.append(translated_text)
        next_idx = i
    
    if chunk_by == "sentences":
        translated_document = reassemble_sentences(translated_chunks, chunk_metadata)
    else:
        translated_document = reassemble_paragraphs(translated_chunks, chunk_metadata)
    
    with open(output_text_file, 'w', encoding='utf-8') as f:
        f.write(translated_document)
    
    return next_idx if translated_chunks else start_idx


def _get_all_runs(paragraph):
    return list(paragraph.runs)


def _join_run_texts(all_runs):
    return ''.join(run.text or '' for run in all_runs)


def _get_run_format_key(run):
    font_color = None
    if run.font.color and run.font.color.rgb:
        font_color = str(run.font.color.rgb)
    # Explicit black is visually identical to inherited (None)
    if font_color == '000000':
        font_color = None
    
    return (
        run.bold,
        run.italic,
        run.underline,
        run.font.name,
        run.font.size,
        font_color
    )


def _merge_identical_runs(paragraph):
    all_runs = _get_all_runs(paragraph)
    if len(all_runs) <= 1:
        return
    
    run_info = []
    for run in all_runs:
        if run.text:
            run_info.append((run, _get_run_format_key(run)))
    
    if len(run_info) <= 1:
        return
    
    i = 0
    while i < len(run_info) - 1:
        current_run, current_key = run_info[i]
        next_run, next_key = run_info[i + 1]
        
        if current_key == next_key:
            current_run.text += next_run.text
            next_run.text = ''
            run_info.pop(i + 1)
        else:
            i += 1


def _build_format_segments(paragraph):
    segments = []
    current_segment_runs = []
    current_format_key = None
    
    for run in _get_all_runs(paragraph):
        if not run.text:
            continue
        
        format_key = _get_run_format_key(run)
        
        if current_format_key is None:
            current_format_key = format_key
            current_segment_runs = [run]
        elif format_key == current_format_key:
            current_segment_runs.append(run)
        else:
            if current_segment_runs:
                segments.append(current_segment_runs)
            current_format_key = format_key
            current_segment_runs = [run]
    
    if current_segment_runs:
        segments.append(current_segment_runs)
    
    return segments


def _has_formatting_differences(paragraph):
    format_keys = set()
    for run in _get_all_runs(paragraph):
        if run.text and run.text.strip():
            format_keys.add(_get_run_format_key(run))
            if len(format_keys) > 1:
                return True
    return False


def _distribute_text_to_runs(translated_text, content_runs, original_lengths):
    total_original = sum(original_lengths)
    split_points = []
    cumulative = 0
    for length in original_lengths[:-1]:
        cumulative += length
        ratio = cumulative / total_original
        raw_offset = int(ratio * len(translated_text))
        # Walk forward to nearest space, then fall back to walking backward
        forward = raw_offset
        while forward < len(translated_text) and translated_text[forward] != ' ':
            forward += 1
        backward = raw_offset
        while backward > 0 and translated_text[backward] != ' ':
            backward -= 1
        if forward < len(translated_text):
            boundary = forward
        else:
            boundary = backward
        split_points.append(boundary)
    
    pieces = []
    previous = 0
    for point in split_points:
        pieces.append(translated_text[previous:point].strip())
        previous = point
    pieces.append(translated_text[previous:].strip())
    
    # Fall back if any content run would be left empty while others have text
    non_empty_pieces = [p for p in pieces if p]
    if len(non_empty_pieces) < len(content_runs) and len(non_empty_pieces) > 0:
        pieces = [translated_text] + [''] * (len(content_runs) - 1)
    
    for run, piece in zip(content_runs, pieces):
        run.text = piece


def _chunk_and_translate(text, translation_manager, source_lang, target_lang, use_find_replace, idx, use_cache=True):
    MAX_CHAR = 600
    if len(text) <= MAX_CHAR:
        result = translation_manager.translate_with_best_model(
            text=text,
            source_lang=source_lang,
            target_lang=target_lang,
            use_find_replace=use_find_replace,
            idx=idx,
            use_cache=use_cache
        )
        return result.get("translated_text", "[TRANSLATION FAILED]")
    
    sentences = _split_into_sentences(text)
    chunks = []
    current_chunk = ''
    for sentence in sentences:
        if len(current_chunk) + len(sentence) + 1 <= MAX_CHAR:
            current_chunk += (' ' if current_chunk else '') + sentence
        else:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = sentence
    if current_chunk:
        chunks.append(current_chunk)
    
    translated_parts = []
    for chunk in chunks:
        result = translation_manager.translate_with_best_model(
            text=chunk,
            source_lang=source_lang,
            target_lang=target_lang,
            use_find_replace=use_find_replace,
            idx=idx,
            use_cache=use_cache
        )
        translated_parts.append(result.get("translated_text", "[TRANSLATION FAILED]"))
    return ' '.join(translated_parts)


def _translate_paragraph(paragraph, translation_manager, source_lang, target_lang, use_find_replace, idx, use_cache=True, hyperlink_records=None):
    from docx.oxml.ns import qn
    
    _apply_cyan = False
    p_elem = paragraph._element
    hyperlink_elems = p_elem.findall(qn('w:hyperlink'))
    has_hyperlinks = len(hyperlink_elems) > 0
    
    if hyperlink_records is not None and has_hyperlinks:
        wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        # Build full paragraph text for context (including hyperlink run text)
        all_text_parts = []
        for child in p_elem:
            if child.tag == qn('w:r'):
                t_elem = child.find(f'{{{wns}}}t')
                if t_elem is not None and t_elem.text:
                    all_text_parts.append(t_elem.text)
            elif child.tag == qn('w:hyperlink'):
                for r_elem in child.findall(f'{{{wns}}}r'):
                    t_elem = r_elem.find(f'{{{wns}}}t')
                    if t_elem is not None and t_elem.text:
                        all_text_parts.append(t_elem.text)
        full_paragraph_text = ''.join(all_text_parts)
        
        # Collect hyperlink records before stripping
        for hl_elem in hyperlink_elems:
            r_id = hl_elem.get(qn('r:id'))
            if r_id and r_id in paragraph.part.rels:
                url = paragraph.part.rels[r_id].target_ref
            else:
                url = ''
            for r_elem in hl_elem.findall(f'{{{wns}}}r'):
                t_elem = r_elem.find(f'{{{wns}}}t')
                original_text = t_elem.text if t_elem is not None and t_elem.text else ''
                hyperlink_records.append({
                    'original_text': original_text,
                    'full_sentence': full_paragraph_text,
                    'url': url,
                })
        
        # Strip hyperlink XML wrappers — move w:r elements up into w:p
        for hl_elem in list(p_elem.findall(qn('w:hyperlink'))):
            for r_elem in list(hl_elem.findall(qn('w:r'))):
                p_elem.insert(list(p_elem).index(hl_elem), r_elem)
            p_elem.remove(hl_elem)
        
        _apply_cyan = True
    
    all_runs = _get_all_runs(paragraph)
    if not all_runs:
        return idx
    
    full_text = _join_run_texts(all_runs)
    
    if not full_text.strip():
        return idx
    
    # Clean up invisible run boundaries by merging identical adjacent runs
    _merge_identical_runs(paragraph)
    
    # Re-fetch after merge since runs may have changed
    all_runs = _get_all_runs(paragraph)
    
    # Check if paragraph has actual formatting differences
    if not _has_formatting_differences(paragraph):
        # No formatting differences - translate as single unit (original simple approach)
        leading_ws = ''
        trailing_ws = ''
        text_to_translate = full_text
        
        if text_to_translate and text_to_translate[0].isspace():
            i = 0
            while i < len(text_to_translate) and text_to_translate[i].isspace():
                i += 1
            leading_ws = text_to_translate[:i]
            text_to_translate = text_to_translate[i:]
        
        if text_to_translate and text_to_translate[-1].isspace():
            i = len(text_to_translate) - 1
            while i >= 0 and text_to_translate[i].isspace():
                i -= 1
            trailing_ws = text_to_translate[i + 1:]
            text_to_translate = text_to_translate[:i + 1]
        
        if not text_to_translate:
            return idx
        
        translated_text = _chunk_and_translate(
            text_to_translate, translation_manager, source_lang, target_lang,
            use_find_replace, idx, use_cache
        )
        translated_text = normalize_apostrophes(translated_text)
        
        first_content_run = None
        for run in all_runs:
            if run.text and run.text.strip():
                first_content_run = run
                break
        
        if first_content_run:
            first_content_run.text = leading_ws + translated_text + trailing_ws
            found_first = False
            for run in all_runs:
                if run.text and run.text.strip():
                    if found_first:
                        run.text = ''
                    else:
                        found_first = True
    
    else:
        # Has formatting differences - translate as single unit and remap proportionally
        content_runs = [run for run in all_runs if run.text and run.text.strip()]
        full_text = _join_run_texts(all_runs)
        original_lengths = [len(run.text) for run in content_runs]
        
        text_to_translate = full_text.strip()
        if not text_to_translate:
            return idx
        
        translated_text = _chunk_and_translate(
            text_to_translate, translation_manager, source_lang, target_lang,
            use_find_replace, idx, use_cache
        )
        translated_text = normalize_apostrophes(translated_text)
        
        _distribute_text_to_runs(translated_text, content_runs, original_lengths)
    
    if _apply_cyan:
        from docx.enum.text import WD_COLOR_INDEX
        for run in _get_all_runs(paragraph):
            if hasattr(run, 'font') and hasattr(run.font, 'highlight_color'):
                run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
    
    return idx + 1


def _set_proofing_language(document, target_lang):
    from docx.oxml.ns import qn
    from lxml import etree
    
    locale_map = {'fr': 'fr-CA', 'en': 'en-CA'}
    locale_code = locale_map[target_lang]
    
    for r_elem in document.element.iter(qn('w:r')):
        rPr = r_elem.find(qn('w:rPr'))
        if rPr is None:
            rPr = etree.SubElement(r_elem, qn('w:rPr'))
            r_elem.insert(0, rPr)
        lang = rPr.find(qn('w:lang'))
        if lang is None:
            lang = etree.SubElement(rPr, qn('w:lang'))
        lang.set(qn('w:val'), locale_code)


def translate_word_document(
        input_docx_file,
        output_docx_file=None,
        source_lang="en",
        models_to_use=None,
        use_find_replace=True,
        use_finetuned=True,
        translation_manager=None,
        include_timestamp=True,
        use_cache=True
):
    import os
    from datetime import datetime
    from docx import Document
    
    if not output_docx_file:
        base, ext = os.path.splitext(input_docx_file)
        date_str = datetime.now().strftime("%Y%m%d")
        if include_timestamp:
            output_docx_file = f"{base}_translated_{date_str}.docx"
        else:
            output_docx_file = f"{base}_translated.docx"
    
    if source_lang not in ["en", "fr"]:
        raise ValueError('source_lang must be either "fr" or "en"')
    
    target_lang = "fr" if source_lang == "en" else "en"
    
    if not translation_manager:
        translation_manager = create_translator(
            use_finetuned=use_finetuned,
            models_to_use=models_to_use,
            use_embedder=True,
            load_models=True
        )
    
    document = Document(input_docx_file)
    idx = 1
    hyperlink_records = []
    
    for paragraph in document.paragraphs:
        idx = _translate_paragraph(paragraph, translation_manager, source_lang, target_lang, use_find_replace, idx, use_cache=use_cache, hyperlink_records=hyperlink_records)
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    idx = _translate_paragraph(paragraph, translation_manager, source_lang, target_lang, use_find_replace, idx, use_cache=use_cache, hyperlink_records=hyperlink_records)
    
    translated_hf_ids = set()
    
    header_footer_attrs = [
        'header', 'footer',
        'first_page_header', 'first_page_footer',
        'even_page_header', 'even_page_footer',
    ]
    
    for section in document.sections:
        for attr in header_footer_attrs:
            hf = getattr(section, attr)
            if id(hf._element) in translated_hf_ids:
                continue
            if hf.is_linked_to_previous:
                continue
            translated_hf_ids.add(id(hf._element))
            for paragraph in hf.paragraphs:
                idx = _translate_paragraph(paragraph, translation_manager, source_lang, target_lang, use_find_replace, idx, use_cache=use_cache, hyperlink_records=hyperlink_records)
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            idx = _translate_paragraph(paragraph, translation_manager, source_lang, target_lang, use_find_replace, idx, use_cache=use_cache, hyperlink_records=hyperlink_records)
    
    _set_proofing_language(document, target_lang)
    document.save(output_docx_file)
    
    if hyperlink_records:
        from translate.hyperlink_notes import write_hyperlink_notes
        notes_path = os.path.splitext(output_docx_file)[0] + '_translation_notes.docx'
        write_hyperlink_notes(hyperlink_records, notes_path)
    
    return output_docx_file
