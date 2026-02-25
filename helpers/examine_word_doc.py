from docx import Document
from pathlib import Path


def list_attr(obj):
    return [x for x in dir(obj) if x[0] != "_"]


def print_dir():
    for group in [document.paragraphs, document.tables, document.sections]:
        for item in group:
            print()
            print(item.__class__.__name__)
            print([x for x in dir(item) if x[0] != "_"])
            if item.__class__.__name__ == "Table":
                for row in item.rows:
                    print('\trows:', list_attr(row))
                    for cell in row.cells:
                        print('\t\tcells:', list_attr(cell))
                        for paragraph in cell.paragraphs:
                            print('\t\t\tcell paragraphs:', list_attr(paragraph))
                            for run in paragraph.runs:
                                print('\t\t\t\tcell paragraph runs:', list_attr(run))
                                break
                            break
                        break
                    break


def print_paragraph_details(p, p_index=None, extra_text=None, print_text=False, min_runs=1, only_include_format_changes=False):
    rows = []
    changes = None
    has_changes = False
    W = 120
    
    idx = 0
    for idx, run in enumerate(p.runs):
        b = "Y" if run.bold else "N"
        it = "Y" if run.italic else "N"
        sz = f"{run.font.size.pt:.1f}" if run.font.size else "Def"
        fnt = str(run.font.name)[:10] if run.font.name else "Def"
        
        color = "Def"
        if run.font.color and run.font.color.rgb:
            color = str(run.font.color.rgb)
            
        if only_include_format_changes and not has_changes:
            if changes is None:
                changes = [b, it, sz, fnt, color]
            elif changes != [b, it, sz, fnt, color]:
                has_changes = True
        
        run_text = run.text.replace("\n", "<br>")
        rows.append(f'{idx:<4} | {b:<4} | {it:<4} | {sz:<5} | {fnt:<10} | {color:<8} | "{run_text}"')
    
    if (not only_include_format_changes or has_changes) and (idx >= min_runs - 1):
        print()
        if p_index is not None and extra_text:
            print_block(f"paragraph {p_index} - {extra_text}", level=2)
            print()
        elif p_index is not None:
            print_block(f"paragraph {p_index}", level=2)
            print()
        elif extra_text:
            print_block(extra_text, level=2)
            print()
            
        print("-" * W)
        print(f"{'Idx':<4} | {'Bold':<4} | {'Ital':<4} | {'Size':<5} | {'Font':<10} | {'Color':<8} | {'Text'}")
        print("-" * W)
        print("\n".join(rows))
        if print_text:
            print()
            print(f'-> full text: "{p.text}"')
            print()
    elif p_index is not None:
        print_block(f"skipping paragraph {p_index} - no formatting changes detected", level=3)


def print_section_details(doc):
    W = 20
    for i, section in enumerate(doc.sections):
        print_block(f"Section {i}", level=3)
        print()
        
        print(f"{'Start Type:':<{W}} {section.start_type}")
        print(f"{'Orientation:':<{W}} {section.orientation}")
        print(f"{'Page Size:':<{W}} {section.page_width.inches:.2f}\" x {section.page_height.inches:.2f}\"")
        print(f"{'Margins (L/R):':<{W}} {section.left_margin.inches:.2f}\" / {section.right_margin.inches:.2f}\"")
        print(f"{'First Page Header:':<{W}} {section.different_first_page_header_footer}")
        
        has_header = any(p.text.strip() for p in section.header.paragraphs)
        print(f"{'Header Text:':<{W}} {has_header}")
        print()


def print_block(text, level=1):
    W = 120
    if level == 0:
        print()
        centered = f"  {text.upper()}  ".center(W, "=")
        print(f"\n{'=' * W}\n{'=' * W}\n{centered}\n{'=' * W}\n{'=' * W}\n")
    elif level == 1:
        print()
        centered = f" {text.upper()} ".center(W, "=")
        print(f"\n{'=' * W}\n{centered}\n{'=' * W}\n")
    elif level == 2:
        print(f" {text} ".center(W, "="))
    elif level == 3:
        print(f"  {text}  ".center(W, "-"))


def print_document_details(document_name, root_directory, display_sections, display_paragraphs, display_tables, **kwargs):
    print_block(document_name, level=0)
    document = Document(root_directory / document_name)
    include_all = not kwargs["only_include_format_changes"]
    
    if display_sections:
        print_block(f"sections ({len(document.sections)})", level=1)
        print_section_details(document)
        
    if display_paragraphs:
        print_block(f"paragraphs ({len(document.paragraphs)})", level=1)
        for i, paragraph in enumerate(document.paragraphs):
            print_paragraph_details(paragraph, p_index=i, **kwargs)
        
    if display_tables:
        print_block(f"tables ({len(document.tables)})", level=1)
        for n, table in enumerate(document.tables):
            print_block(f"table {n}", level=2)
            for i, row in enumerate(table.rows):
                if include_all:
                    print_block(f"row {i}", level=3)
                for j, cell in enumerate(row.cells):
                    if include_all:
                        print_block(f"cell {j}", level=3)
                    for k, paragraph in enumerate(cell.paragraphs):
                        extra_text = None if include_all else f" row, col = {i}, {j} "
                        print_paragraph_details(paragraph, extra_text=extra_text, **kwargs)
                if include_all:
                    print()


if __name__ == '__main__':
    kwargs = {
        "print_text": True,
        "min_runs": 1,
        "only_include_format_changes": True,
    }
    root = Path('..') / "_TRANSLATED_DOCUMENTS" / "fsar_docs"
    docs = [
        "1301_en.docx",
        # "1303_en.docx",
        # "1407_en.docx",
        # "1429_en.docx",
        # "1379_fr.docx",
        # "1415_fr.docx",
        # "1444_fr.docx",
        # "1458_fr.docx",
    ]

    for doc in docs:
        print_document_details(
            doc,
            root,
            display_sections=True,
            display_paragraphs=True,
            display_tables=True,
            **kwargs
        )
        