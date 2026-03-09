from unittest.mock import patch
from docx import Document
from scitrans.translate.word_document import _translate_table_cell, write_translations_notes


class MockTranslator:
    def __init__(self):
        self.call_count = 0
        self.source_texts = []
    
    def translate_with_best_model(self, text, source_lang, target_lang, use_find_replace, idx, **kwargs):
        self.call_count += 1
        self.source_texts.append(text)
        return {"translated_text": f"[TR:{text}]"}


def make_cell(text):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    return cell


def make_cell_with_runs(run_texts):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    # Clear default paragraph
    cell.paragraphs[0].clear()
    for i, txt in enumerate(run_texts):
        run = cell.paragraphs[0].add_run(txt)  # FIXME???
    return cell


def make_cell_with_paragraphs(para_texts):
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.paragraphs[0].text = para_texts[0]
    for txt in para_texts[1:]:
        cell.add_paragraph(txt)
    return cell


# ---------------------------------------------------------------------------
# 1. Numeric conversion
# ---------------------------------------------------------------------------
class TestNumericConversion:
    def _enabled_config(self):
        return {"enabled": True}
    
    def _disabled_config(self):
        return {"enabled": False}
    
    @patch("scitrans.translate.word_document.config")
    def test_integer_en_to_fr(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = self._enabled_config()
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        cell = make_cell("1234")
        idx = _translate_table_cell(cell, None, "en", "fr", False, 0)
        assert cell.text.replace("\xa0", " ") == "1 234"
        assert idx == 0
    
    @patch("scitrans.translate.word_document.config")
    def test_decimal_en_to_fr(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = self._enabled_config()
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        cell = make_cell("1,234.56")
        _translate_table_cell(cell, None, "en", "fr", False, 0)
        result = cell.text.replace("\xa0", " ")
        assert result == "1 234,56"
    
    @patch("scitrans.translate.word_document.config")
    def test_percentage_en_to_fr(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = self._enabled_config()
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        cell = make_cell("45.6%")
        _translate_table_cell(cell, None, "en", "fr", False, 0)
        result = cell.text.replace("\xa0", " ")
        assert result == "45,6 %"
    
    @patch("scitrans.translate.word_document.config")
    def test_integer_fr_to_en(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = self._enabled_config()
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        cell = make_cell("1234")
        _translate_table_cell(cell, None, "fr", "en", False, 0)
        assert cell.text == "1,234"
    
    @patch("scitrans.translate.word_document.config")
    def test_decimal_fr_to_en(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = self._enabled_config()
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        cell = make_cell("1 234,56".replace(" ", "\xa0"))
        _translate_table_cell(cell, None, "fr", "en", False, 0)
        assert cell.text == "1,234.56"
    
    @patch("scitrans.translate.word_document.config")
    def test_multi_run_writes_first_clears_rest(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = self._enabled_config()
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        cell = make_cell_with_runs(["12", "34"])
        _translate_table_cell(cell, None, "en", "fr", False, 0)
        runs = cell.paragraphs[0].runs
        # First non-empty run gets the converted value, rest cleared
        non_empty = [r.text for r in runs if r.text]
        assert len(non_empty) <= 1
    
    @patch("scitrans.translate.word_document.config")
    def test_returns_idx_unchanged(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = self._enabled_config()
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        cell = make_cell("999")
        idx = _translate_table_cell(cell, None, "en", "fr", False, 42)
        assert idx == 42
    
    @patch("scitrans.translate.word_document.config")
    def test_disabled_falls_through(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = self._disabled_config()
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("1234")
        # With numeric disabled and high AI threshold, short numeric text left as-is
        _translate_table_cell(cell, None, "en", "fr", False, 0)
        assert cell.text == "1234"
    
    @patch("scitrans.translate.word_document.config")
    def test_whitespace_padded_number(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = self._enabled_config()
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        cell = make_cell("  5678  ")
        _translate_table_cell(cell, None, "en", "fr", False, 0)
        result = cell.text.replace("\xa0", " ")
        assert "5 678" in result


# ---------------------------------------------------------------------------
# 2. Table translations dict
# ---------------------------------------------------------------------------
class TestTableTranslationsDict:
    @patch("scitrans.translate.word_document.config")
    def test_exact_match_replaces(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("Yes")
        idx = _translate_table_cell(cell, None, "en", "fr", False, 0,
                                    table_translations_dict={"Yes": "Oui"})
        assert cell.text == "Oui"
        assert idx == 0
    
    @patch("scitrans.translate.word_document.config")
    def test_case_sensitive_no_match(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("yes")
        _translate_table_cell(cell, None, "en", "fr", False, 0,
                              table_translations_dict={"Yes": "Oui"})
        assert cell.text == "yes"
    
    @patch("scitrans.translate.word_document.config")
    def test_missing_key_falls_through(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("Maybe")
        _translate_table_cell(cell, None, "en", "fr", False, 0,
                              table_translations_dict={"Yes": "Oui"})
        assert cell.text == "Maybe"
    
    @patch("scitrans.translate.word_document.config")
    def test_none_dict_falls_through(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("Yes")
        _translate_table_cell(cell, None, "en", "fr", False, 0,
                              table_translations_dict=None)
        assert cell.text == "Yes"
    
    @patch("scitrans.translate.word_document.config")
    def test_empty_dict_falls_through(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("Yes")
        _translate_table_cell(cell, None, "en", "fr", False, 0,
                              table_translations_dict={})
        assert cell.text == "Yes"
    
    @patch("scitrans.translate.word_document.config")
    def test_multi_run_first_gets_value_rest_cleared(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell_with_runs(["Y", "es"])
        _translate_table_cell(cell, None, "en", "fr", False, 0,
                              table_translations_dict={"Yes": "Oui"})
        runs = cell.paragraphs[0].runs
        texts = [r.text for r in runs]
        assert "".join(texts) == "Oui"
    
    @patch("scitrans.translate.word_document.config")
    def test_returns_idx_unchanged(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("Yes")
        idx = _translate_table_cell(cell, None, "en", "fr", False, 7,
                                    table_translations_dict={"Yes": "Oui"})
        assert idx == 7


# ---------------------------------------------------------------------------
# 3. Preferential translations
# ---------------------------------------------------------------------------
class TestPreferentialTranslations:
    def _pref_dict(self):
        return {
            "translations": {
                "general": {
                    "Environnement": {"en": "Environment"},
                    "Analyse": {"en": "Analysis"},
                }
            }
        }
    
    @patch("scitrans.translate.word_document.config")
    def test_case_insensitive_match_fr_to_en(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("environnement")
        _translate_table_cell(cell, None, "fr", "en", False, 0,
                              preferential_dict=self._pref_dict())
        # FR→EN: match_translation = term_key = "Environnement"
        assert cell.text == "Environnement"
    
    @patch("scitrans.translate.word_document.config")
    def test_capitalization_preserved(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("Environnement")
        _translate_table_cell(cell, None, "fr", "en", False, 0,
                              preferential_dict=self._pref_dict())
        assert cell.text == "Environnement"
    
    @patch("scitrans.translate.word_document.config")
    def test_en_to_fr_uses_get_translation_value(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        # EN→FR: term_key.lower() must match cell text (lowered)
        # term_key is "Environnement" (French), cell text is "Environment" (English)
        # These won't match since the lookup compares term_key to stripped text
        # Build a dict where term_key matches the English text
        pref = {
            "translations": {
                "general": {
                    "Environment": {"en": "Environnement"},
                }
            }
        }
        cell = make_cell("Environment")
        _translate_table_cell(cell, None, "en", "fr", False, 0,
                              preferential_dict=pref)
        # get_translation_value returns term_data.get('en') = "Environnement"
        assert cell.text == "Environnement"
    
    @patch("scitrans.translate.word_document.config")
    def test_no_translations_wrapper(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        # Dict without "translations" key — falls back to using dict directly
        pref = {
            "general": {
                "Environnement": {"en": "Environment"},
            }
        }
        cell = make_cell("environnement")
        _translate_table_cell(cell, None, "fr", "en", False, 0,
                              preferential_dict=pref)
        assert cell.text == "Environnement"
    
    @patch("scitrans.translate.word_document.config")
    def test_no_match_falls_through(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("Unknown")
        _translate_table_cell(cell, None, "en", "fr", False, 0,
                              preferential_dict=self._pref_dict())
        assert cell.text == "Unknown"
    
    @patch("scitrans.translate.word_document.config")
    def test_returns_idx_unchanged(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("environnement")
        idx = _translate_table_cell(cell, None, "fr", "en", False, 5,
                                    preferential_dict=self._pref_dict())
        assert idx == 5
    
    @patch("scitrans.translate.word_document.config")
    def test_multi_run_first_gets_value(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell_with_runs(["Environ", "nement"])
        _translate_table_cell(cell, None, "fr", "en", False, 0,
                              preferential_dict=self._pref_dict())
        combined = "".join(r.text for r in cell.paragraphs[0].runs)
        assert combined == "Environnement"


# ---------------------------------------------------------------------------
# 4. AI translation (long text)
# ---------------------------------------------------------------------------
class TestAITranslation:
    @patch("scitrans.translate.word_document._translate_paragraph")
    @patch("scitrans.translate.word_document.config")
    def test_long_text_triggers_translate_paragraph(self, mock_config, mock_tp):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        mock_tp.return_value = 1
        cell = make_cell("This is a long sentence for AI translation")
        idx = _translate_table_cell(cell, "mgr", "en", "fr", False, 0)
        assert mock_tp.called
    
    @patch("scitrans.translate.word_document._translate_paragraph")
    @patch("scitrans.translate.word_document.config")
    def test_exactly_at_threshold_triggers(self, mock_config, mock_tp):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        mock_tp.return_value = 1
        cell = make_cell("a" * 20)
        _translate_table_cell(cell, "mgr", "en", "fr", False, 0)
        assert mock_tp.called
    
    @patch("scitrans.translate.word_document._translate_paragraph")
    @patch("scitrans.translate.word_document.config")
    def test_one_below_threshold_does_not_trigger(self, mock_config, mock_tp):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        mock_tp.return_value = 1
        cell = make_cell("a" * 19)
        _translate_table_cell(cell, "mgr", "en", "fr", False, 0)
        assert not mock_tp.called
    
    @patch("scitrans.translate.word_document._translate_paragraph")
    @patch("scitrans.translate.word_document.config")
    def test_multi_paragraph_each_translated(self, mock_config, mock_tp):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        mock_tp.side_effect = lambda *a, **kw: a[5] + 1  # increment idx
        cell = make_cell_with_paragraphs(["First paragraph that is long enough", "Second paragraph also long enough"])
        idx = _translate_table_cell(cell, "mgr", "en", "fr", False, 0)
        assert mock_tp.call_count == 2
        assert idx == 2
    
    @patch("scitrans.translate.word_document._translate_paragraph")
    @patch("scitrans.translate.word_document.config")
    def test_idx_incremented_per_paragraph(self, mock_config, mock_tp):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 5}
        mock_tp.side_effect = lambda *a, **kw: a[5] + 1
        cell = make_cell_with_paragraphs(["Hello world", "Goodbye world", "Third line"])
        idx = _translate_table_cell(cell, "mgr", "en", "fr", False, 10)
        assert idx == 13
    
    @patch("scitrans.translate.word_document._translate_paragraph")
    @patch("scitrans.translate.word_document.config")
    def test_custom_threshold(self, mock_config, mock_tp):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 50}
        mock_tp.return_value = 1
        cell = make_cell("a" * 49)
        _translate_table_cell(cell, "mgr", "en", "fr", False, 0)
        assert not mock_tp.called
        cell2 = make_cell("a" * 50)
        _translate_table_cell(cell2, "mgr", "en", "fr", False, 0)
        assert mock_tp.called


# ---------------------------------------------------------------------------
# 5. Short text left as-is
# ---------------------------------------------------------------------------
class TestShortTextLeftAsIs:
    @patch("scitrans.translate.word_document.config")
    def test_short_text_unchanged(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("Hello")
        _translate_table_cell(cell, None, "en", "fr", False, 0)
        assert cell.text == "Hello"
    
    @patch("scitrans.translate.word_document.config")
    def test_not_numeric_not_in_dicts(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": True}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("abc")
        _translate_table_cell(cell, None, "en", "fr", False, 0)
        assert cell.text == "abc"
    
    @patch("scitrans.translate.word_document.config")
    def test_returns_same_idx(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("Hello")
        idx = _translate_table_cell(cell, None, "en", "fr", False, 99)
        assert idx == 99
    
    @patch("scitrans.translate.word_document.config")
    def test_short_no_dicts_provided(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("Word")
        idx = _translate_table_cell(cell, None, "en", "fr", False, 0,
                                    preferential_dict=None, table_translations_dict=None)
        assert cell.text == "Word"
        assert idx == 0


# ---------------------------------------------------------------------------
# 6. Edge cases
# ---------------------------------------------------------------------------
class TestEdgeCases:
    @patch("scitrans.translate.word_document.config")
    def test_empty_cell(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": True}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        cell = make_cell("")
        idx = _translate_table_cell(cell, None, "en", "fr", False, 5)
        assert idx == 5
    
    @patch("scitrans.translate.word_document.config")
    def test_whitespace_only(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": True}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 20}
        cell = make_cell("   ")
        idx = _translate_table_cell(cell, None, "en", "fr", False, 5)
        assert idx == 5
    
    @patch("scitrans.translate.word_document.config")
    def test_leading_trailing_whitespace_stripped_for_match(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("  Yes  ")
        _translate_table_cell(cell, None, "en", "fr", False, 0,
                              table_translations_dict={"Yes": "Oui"})
        assert "Oui" in cell.text


# ---------------------------------------------------------------------------
# 7. Dispatch priority
# ---------------------------------------------------------------------------
class TestDispatchPriority:
    @patch("scitrans.translate.word_document.config")
    def test_numeric_beats_dict(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": True}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        cell = make_cell("1234")
        _translate_table_cell(cell, None, "en", "fr", False, 0,
                              table_translations_dict={"1234": "one-two-three-four"})
        # Numeric fires first, not the dict
        assert cell.text != "one-two-three-four"
    
    @patch("scitrans.translate.word_document.config")
    def test_dict_beats_preferential(self, mock_config):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 200}
        pref = {"translations": {"general": {"Yes": {"en": "PrefYes"}}}}
        cell = make_cell("Yes")
        _translate_table_cell(cell, None, "en", "fr", False, 0,
                              table_translations_dict={"Yes": "DictOui"},
                              preferential_dict=pref)
        assert cell.text == "DictOui"
    
    @patch("scitrans.translate.word_document._translate_paragraph")
    @patch("scitrans.translate.word_document.config")
    def test_preferential_beats_ai(self, mock_config, mock_tp):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 5}
        mock_tp.return_value = 1
        pref = {"translations": {"general": {"environnement": {"en": "Environment"}}}}
        cell = make_cell("environnement")
        _translate_table_cell(cell, "mgr", "fr", "en", False, 0,
                              preferential_dict=pref)
        # Preferential should fire, not AI
        assert not mock_tp.called


# ---------------------------------------------------------------------------
# 8. Formatting rules (via _translate_paragraph path)
# ---------------------------------------------------------------------------
class TestFormattingRulesInTables:
    @patch("scitrans.translate.word_document._translate_paragraph")
    @patch("scitrans.translate.word_document.config")
    def test_long_cell_goes_through_paragraph_translation(self, mock_config, mock_tp):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 10}
        mock_tp.return_value = 1
        cell = make_cell("Some long text that needs formatting rules applied")
        _translate_table_cell(cell, "mgr", "en", "fr", False, 0)
        # _translate_paragraph is where formatting rules fire
        assert mock_tp.called
    
    @patch("scitrans.translate.word_document._translate_paragraph")
    @patch("scitrans.translate.word_document.config")
    def test_short_cell_skips_paragraph_translation(self, mock_config, mock_tp):
        mock_config.NUMERIC_CONVERSION_CONFIG = {"enabled": False}
        mock_config.TABLE_TRANSLATION_CONFIG = {"min_cell_length_for_ai": 100}
        mock_tp.return_value = 1
        cell = make_cell("Short")
        _translate_table_cell(cell, "mgr", "en", "fr", False, 0)
        assert not mock_tp.called


# ---------------------------------------------------------------------------
# 9. Hyperlink notes
# ---------------------------------------------------------------------------
class TestHyperlinkNotes:
    def test_row_count_matches_records(self, tmp_path):
        records = [
            {"original_text": "a", "full_paragraph": "b", "notes": "c"},
            {"original_text": "d", "full_paragraph": "e", "notes": "f"},
            {"original_text": "g", "full_paragraph": "h", "notes": "i"},
        ]
        out = str(tmp_path / "notes.docx")
        write_translations_notes(records, out)
        doc = Document(out)
        table = doc.tables[0]
        # 1 header row + 3 data rows
        assert len(table.rows) == 4
