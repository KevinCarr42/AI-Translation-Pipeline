import pytest
from pathlib import Path
from docx import Document
from scitrans.translate.models import create_translator
from scitrans.translate.document import translate_word_document


@pytest.fixture(params=[("en", "fr"), ("fr", "en")])
def translation_context(request, tmp_path):
    source_lang, target_lang = request.param
    base_path = Path(__file__).parent / "fixtures"
    
    input_file = base_path / f"test_table_replacements_{source_lang}.docx"
    goal_file = base_path / f"test_table_replacements_{target_lang}.docx"
    output_file = tmp_path / f"translated_{source_lang}_to_{target_lang}.docx"
    
    translate_word_document(
        input_docx_file=str(input_file),
        output_docx_file=str(output_file),
        source_lang=source_lang,
        models_to_use=None,
        use_find_replace=True,
        translation_manager=create_translator(),
        include_timestamp=False,
        use_cache=True
    )
    
    return {
        "doc": Document(output_file),
        "doc_goal": Document(goal_file),
        "direction": f"{source_lang}->{target_lang}"
    }


class TestDocumentTranslations:
    def test_paragraphs(self, translation_context, subtests):
        doc = translation_context["doc"]
        goal = translation_context["doc_goal"]
        
        for i, (p, p_goal) in enumerate(zip(doc.paragraphs, goal.paragraphs)):
            with subtests.test(msg=f"Paragraph {i}", direction=translation_context["direction"]):
                assert p.text == p_goal.text
    
    def test_tables(self, translation_context, subtests):
        doc = translation_context["doc"]
        goal = translation_context["doc_goal"]
        
        for t_idx, (t, t_goal) in enumerate(zip(doc.tables, goal.tables)):
            for r_idx, (r, r_goal) in enumerate(zip(t.rows, t_goal.rows)):
                for c_idx, (c, c_goal) in enumerate(zip(r.cells, r_goal.cells)):
                    with subtests.test(msg="Cell", table=t_idx, row=r_idx, col=c_idx):
                        assert c.text == c_goal.text
