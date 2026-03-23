import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree

from scitrans.translate.word_formatting import (
    SuperscriptOrdinalsRule,
    _split_run_for_vertical_align,
)
from scitrans.translate.word_document import (
    _extract_non_run_elements,
    _reinsert_non_run_elements,
    _convert_newlines_to_breaks,
    _translate_paragraph,
)
from tests.conftest import MockTranslator


# ---------------------------------------------------------------------------
# Bug 1: _split_run_for_vertical_align duplicates text when run has newlines
# ---------------------------------------------------------------------------

class TestSplitRunMultipleWt:
    """When paragraph.text is set and the text contains newlines, python-docx
    creates <w:br/> + multiple <w:t> children inside a single <w:r>.
    _split_run_for_vertical_align deepcopies the run element and only
    updates the first <w:t>, leaving the remaining <w:t> elements as
    duplicated ghost text."""
    
    def test_split_run_preserves_text_with_newlines(self):
        # Simulates the state after paragraph.text = translated_text
        # when the translated text contains newlines and an ordinal
        doc = Document()
        para = doc.add_paragraph("before text\nmiddle 25e after\nlast line")
        
        original_text = para.text
        run = para.runs[0]
        
        # Split at the "e" after "25" (position 21 in the virtual text)
        _split_run_for_vertical_align(para, run, "e", "superscript", offset=21)
        
        # The full paragraph text must be identical — no duplication
        assert para.text == original_text, (
            f"Text was corrupted by split.\n"
            f"  Expected: {original_text!r}\n"
            f"  Got:      {para.text!r}"
        )
    
    def test_superscript_ordinals_no_duplication_with_newlines(self):
        doc = Document()
        text = (
            "Stock reference\n"
            "(USR: 0.8 x BMSY-proxy), (D) Estimated\n"
            "(5-11 mm) from SISCALS (box: 25e, 50e et 75e percentiles)."
        )
        para = doc.add_paragraph(text)
        original_text = para.text
        
        rule = SuperscriptOrdinalsRule()
        rule.apply(para, ["25th", "50th", "75th"])
        
        assert para.text == original_text, (
            f"Ordinal formatting duplicated text.\n"
            f"  Expected length: {len(original_text)}\n"
            f"  Got length:      {len(para.text)}\n"
            f"  Got: {para.text!r}"
        )
    
    def test_superscript_run_has_single_wt(self):
        # The superscript "e" run should contain only a single <w:t> with "e"
        doc = Document()
        para = doc.add_paragraph("line1\nline2 25e end\nline3")
        
        rule = SuperscriptOrdinalsRule()
        rule.apply(para, ["25th"])
        
        for run in para.runs:
            if run.font.superscript:
                t_elems = run._element.findall(qn("w:t"))
                br_elems = run._element.findall(qn("w:br"))
                assert len(t_elems) == 1, (
                    f"Superscript run should have 1 <w:t>, got {len(t_elems)}. "
                    f"Texts: {[t.text for t in t_elems]}"
                )
                assert len(br_elems) == 0, (
                    f"Superscript run should have 0 <w:br>, got {len(br_elems)}"
                )
                assert run.text == "e"
    
    def test_split_then_newline_conversion_no_duplication(self):
        # Full pipeline: ordinal split followed by newline conversion
        doc = Document()
        text = "Caption\n(USR: 0.8), (D)\n(5-11 mm) (box: 25e, 50e, 75e pct)."
        para = doc.add_paragraph(text)
        original_text = para.text
        
        rule = SuperscriptOrdinalsRule()
        rule.apply(para, ["25th", "50th", "75th"])
        _convert_newlines_to_breaks(para)
        
        assert para.text == original_text, (
            f"Text corrupted after split + newline conversion.\n"
            f"  Expected: {original_text!r}\n"
            f"  Got:      {para.text!r}"
        )


class TestTranslateParagraphNewlineOrdinals:
    """End-to-end: _translate_paragraph with text that has newlines and
    ordinals should not duplicate content."""
    
    def test_no_text_duplication(self):
        doc = Document()
        # Source paragraph with superscript ordinals
        para = doc.add_paragraph()
        run1 = para.add_run("Data for 25")
        run2 = para.add_run("th")
        run2.font.superscript = True
        run3 = para.add_run(" and 50")
        run4 = para.add_run("th")
        run4.font.superscript = True
        run5 = para.add_run(" percentile values.")
        
        source_text = para.text
        mock = MockTranslator()
        # MockTranslator wraps text as "[TR:text]", preserving the structure
        # The translated text will contain the ordinals
        
        _translate_paragraph(
            paragraph=para,
            translation_manager=mock,
            source_lang="en",
            target_lang="fr",
            use_find_replace=False,
            idx=1,
            use_cache=False,
            formatting_records=[],
            preferential_dict=None,
            chunk_by="sentences",
            location=None,
        )
        
        translated_text = para.text
        # The mock wraps as "[TR:...]" — text should appear exactly once
        assert translated_text.count("[TR:") == 1, (
            f"Expected exactly 1 [TR:] marker, found {translated_text.count('[TR:')}\n"
            f"  Full text: {translated_text!r}"
        )


# ---------------------------------------------------------------------------
# Bug 2: smartTag elements lose position and their text is invisible
# ---------------------------------------------------------------------------

def _make_paragraph_with_smarttag(doc):
    """Create a paragraph: <w:r>ISSN </w:r><w:smartTag><w:smartTagPr/><w:r>1919-5087</w:r></w:smartTag><w:r> rest</w:r>"""
    para = doc.add_paragraph()
    
    # "ISSN " run
    run1 = para.add_run("ISSN ")
    
    # smartTag wrapping "1919-5087" (with smartTagPr, as in real documents)
    smart_tag = OxmlElement("w:smartTag")
    smart_tag.set(qn("w:uri"), "urn:schemas-microsoft-com:office:smarttags")
    smart_tag.set(qn("w:element"), "phone")
    smart_tag_pr = OxmlElement("w:smartTagPr")
    attr = OxmlElement("w:attr")
    attr.set(qn("w:name"), "phonenumber")
    attr.set(qn("w:val"), "19195087")
    smart_tag_pr.append(attr)
    smart_tag.append(smart_tag_pr)
    inner_r = OxmlElement("w:r")
    inner_t = OxmlElement("w:t")
    inner_t.text = "1919-5087"
    inner_r.append(inner_t)
    smart_tag.append(inner_r)
    para._element.append(smart_tag)
    
    # " rest" run
    run2 = para.add_run("\nISBN 978-0-660-78353-6")
    
    return para


class TestSmartTagHandling:
    """smartTag elements wrap runs that are invisible to paragraph.text and
    paragraph.runs. _extract_non_run_elements removes them and
    _reinsert_non_run_elements appends at the end, losing position."""
    
    def test_smarttag_text_included_in_paragraph_text(self):
        # After extraction+reinsertion, the smartTag text should be
        # part of the normal text flow, not hidden
        doc = Document()
        para = _make_paragraph_with_smarttag(doc)
        
        saved = _extract_non_run_elements(para)
        # The smartTag text should be unwrapped into a regular run
        # so it's visible in paragraph.text
        text_after_extract = para.text
        assert "1919-5087" in text_after_extract, (
            f"smartTag text '1919-5087' should be visible in paragraph.text "
            f"after extraction, got: {text_after_extract!r}"
        )
    
    def test_smarttag_unwrap_no_orphaned_elements(self):
        # smartTagPr must not leak into the paragraph as a direct child
        doc = Document()
        para = _make_paragraph_with_smarttag(doc)
        
        _extract_non_run_elements(para)
        
        for child in para._element:
            assert "smartTagPr" not in child.tag, (
                f"Orphaned {child.tag} found as direct child of paragraph — "
                f"this causes Word to report unreadable content"
            )
    
    def test_smarttag_text_position_preserved(self):
        # The smartTag text should remain between "ISSN " and "\nISBN..."
        doc = Document()
        para = _make_paragraph_with_smarttag(doc)
        
        saved = _extract_non_run_elements(para)
        _reinsert_non_run_elements(para, saved)
        
        full_text = para.text
        issn_pos = full_text.find("ISSN ")
        number_pos = full_text.find("1919-5087")
        isbn_pos = full_text.find("ISBN")
        
        assert number_pos != -1, (
            f"'1919-5087' not found in text: {full_text!r}"
        )
        assert issn_pos < number_pos < isbn_pos, (
            f"Expected 'ISSN ' < '1919-5087' < 'ISBN' but positions are "
            f"ISSN={issn_pos}, 1919-5087={number_pos}, ISBN={isbn_pos}\n"
            f"  Full text: {full_text!r}"
        )
    
    def test_translate_paragraph_preserves_smarttag_text(self):
        # End-to-end: translating a paragraph with a smartTag should
        # keep the smartTag text in position
        doc = Document()
        para = _make_paragraph_with_smarttag(doc)
        mock = MockTranslator()
        
        _translate_paragraph(
            paragraph=para,
            translation_manager=mock,
            source_lang="en",
            target_lang="fr",
            use_find_replace=False,
            idx=1,
            use_cache=False,
            formatting_records=[],
            preferential_dict=None,
            chunk_by="sentences",
            location=None,
        )
        
        translated_text = para.text
        # The number "1919-5087" should appear in the translated text
        assert "1919-5087" in translated_text, (
            f"smartTag number '1919-5087' missing from translated text: "
            f"{translated_text!r}"
        )
        # It should appear before "ISBN"
        num_pos = translated_text.find("1919-5087")
        isbn_pos = translated_text.find("ISBN")
        if isbn_pos != -1:
            assert num_pos < isbn_pos, (
                f"'1919-5087' should come before 'ISBN' but positions are "
                f"1919-5087={num_pos}, ISBN={isbn_pos}\n"
                f"  Full text: {translated_text!r}"
            )


# ---------------------------------------------------------------------------
# Bug 3: orphaned field-end run with blue color contaminates translated text
# ---------------------------------------------------------------------------

def _make_paragraph_with_orphaned_field_end(doc):
    """Create a paragraph that starts with a zero-text blue field-end run,
    followed by normal text runs with a tab — matching the structure found
    in 1432_en.docx paragraph 142."""
    para = doc.add_paragraph()
    p_elem = para._element
    # Remove the default empty run that add_paragraph creates
    for child in list(p_elem):
        if child.tag == qn('w:r'):
            p_elem.remove(child)
    
    # Run 0: orphaned field end — blue color, no w:t, just w:fldChar end
    r0 = OxmlElement('w:r')
    rpr0 = OxmlElement('w:rPr')
    color0 = OxmlElement('w:color')
    color0.set(qn('w:val'), '2B579A')
    rpr0.append(color0)
    r0.append(rpr0)
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'end')
    r0.append(fld_char)
    p_elem.append(r0)
    
    # Run 1: "ISSN 1919-5087" — normal (no color)
    r1 = OxmlElement('w:r')
    t1 = OxmlElement('w:t')
    t1.text = 'ISSN 1919-5087'
    t1.set(qn('xml:space'), 'preserve')
    r1.append(t1)
    p_elem.append(r1)
    
    # Run 2: tab + "Cat. No. XYZ" — normal
    r2 = OxmlElement('w:r')
    r2.append(OxmlElement('w:tab'))
    t2 = OxmlElement('w:t')
    t2.text = 'Cat. No. XYZ'
    r2.append(t2)
    p_elem.append(r2)
    
    return para


class TestOrphanedFieldEndColorContamination:
    """A zero-text run containing w:fldChar end with blue color (2B579A)
    sits at the start of a paragraph. During translation, this run becomes
    group[0] and the translated text is inserted into it, inheriting the
    blue color. The field-end run should be removed before translation so
    its formatting does not contaminate the translated text."""
    
    def test_translated_text_not_blue(self):
        doc = Document()
        para = _make_paragraph_with_orphaned_field_end(doc)
        mock = MockTranslator()
        
        _translate_paragraph(
            paragraph=para,
            translation_manager=mock,
            source_lang='en',
            target_lang='fr',
            use_find_replace=False,
            idx=1,
            use_cache=False,
            formatting_records=[],
            preferential_dict=None,
            chunk_by='sentences',
            location=None,
        )
        
        # No run with visible text should have the blue color
        for run in para.runs:
            if run.text and run.text.strip():
                color_rgb = None
                if run.font.color and run.font.color.rgb:
                    color_rgb = str(run.font.color.rgb)
                assert color_rgb != '2B579A', (
                    f"Run with text {run.text!r} inherited blue color 2B579A "
                    f"from the orphaned field-end run"
                )


# ---------------------------------------------------------------------------
# Bug 4: paragraph.text = ... destroys run-level formatting (font size, bold)
# ---------------------------------------------------------------------------

class TestParagraphTextAssignmentStripsFormatting:
    """When a paragraph has no tabs, _translate_paragraph uses
    paragraph.text = translated_text, which creates a new run with default
    formatting and destroys the original runs' font size, bold, etc.
    This causes table cells with 8.5pt bold text to revert to default size."""
    
    def test_run_font_size_preserved_after_translation(self):
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run('Exploitation Rate Index')
        run.font.size = Pt(8.5)
        run.bold = True
        
        mock = MockTranslator()
        
        _translate_paragraph(
            paragraph=para,
            translation_manager=mock,
            source_lang='en',
            target_lang='fr',
            use_find_replace=False,
            idx=1,
            use_cache=False,
            formatting_records=[],
            preferential_dict=None,
            chunk_by='sentences',
            location=None,
        )
        
        runs_with_text = [r for r in para.runs if r.text and r.text.strip()]
        assert runs_with_text, "Paragraph should have at least one run with text"
        for run in runs_with_text:
            assert run.font.size == Pt(8.5), (
                f"Run with text {run.text!r} lost font size: "
                f"expected {Pt(8.5)}, got {run.font.size}"
            )
    
    def test_multi_run_same_format_preserves_size(self):
        doc = Document()
        para = doc.add_paragraph()
        for text in ['2024 NSAR', ' FB spatiotemporal index (kt):']:
            run = para.add_run(text)
            run.font.size = Pt(8.5)
            run.bold = True
        
        mock = MockTranslator()
        
        _translate_paragraph(
            paragraph=para,
            translation_manager=mock,
            source_lang='en',
            target_lang='fr',
            use_find_replace=False,
            idx=1,
            use_cache=False,
            formatting_records=[],
            preferential_dict=None,
            chunk_by='sentences',
            location=None,
        )
        
        runs_with_text = [r for r in para.runs if r.text and r.text.strip()]
        assert runs_with_text, "Paragraph should have at least one run with text"
        for run in runs_with_text:
            assert run.font.size == Pt(8.5), (
                f"Run with text {run.text!r} lost font size: "
                f"expected {Pt(8.5)}, got {run.font.size}"
            )
