import json
import os
import re

import pytest
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
import docx.oxml.ns as ns

from scitrans.translate.word_document import translate_word_document, _translate_paragraph
from scitrans.translate.word_notes import write_notes_json
from scitrans.translate.utils import split_by_sentences
from scitrans.translate.models import create_translator
from tests.conftest import MockTranslator


# ---------------------------------------------------------------------------
# Helpers for hyperlink tests
# ---------------------------------------------------------------------------

def _add_hyperlink(paragraph, url, link_text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(ns.qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(ns.qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.set(ns.qn('xml:space'), 'preserve')
    t.text = link_text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def _get_full_paragraph_text(paragraph):
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    text_parts = []
    for child in paragraph._element:
        if child.tag == ns.qn('w:r'):
            for t in child.findall('.//w:t', nsmap):
                if t.text:
                    text_parts.append(t.text)
        elif child.tag == ns.qn('w:hyperlink'):
            for t in child.findall('.//w:t', nsmap):
                if t.text:
                    text_parts.append(t.text)
    return ''.join(text_parts)


# ---------------------------------------------------------------------------
# Integration tests (require real ML models)
# ---------------------------------------------------------------------------

@pytest.fixture(scope="module")
def real_translator():
    return create_translator(
        use_finetuned=False,
        models_to_use=['facebook/mbart-large-50-many-to-many-mmt'],
        use_embedder=False,
        load_models=True
    )


@pytest.fixture
def fixture_path():
    return os.path.join(os.path.dirname(__file__), 'fixtures', 'test_document_formatting_en.docx')


@pytest.mark.slow
def test_basic_translation(real_translator, fixture_path, tmp_path):
    output_path = str(tmp_path / 'output.docx')
    translate_word_document(
        input_docx_file=fixture_path,
        output_docx_file=output_path,
        source_lang='en',
        use_find_replace=False,
        translation_manager=real_translator
    )
    
    assert os.path.exists(output_path)
    doc = Document(output_path)
    assert len(doc.paragraphs) > 0


@pytest.mark.slow
def test_table_translation(real_translator, fixture_path, tmp_path):
    output_path = str(tmp_path / 'output.docx')
    translate_word_document(
        input_docx_file=fixture_path,
        output_docx_file=output_path,
        source_lang='en',
        use_find_replace=False,
        translation_manager=real_translator
    )
    
    doc = Document(output_path)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3
    assert len(table.rows[0].cells) == 2
    assert all(
        cell.text.strip() != ''
        for row in table.rows for cell in row.cells
    )


@pytest.mark.slow
def test_empty_runs_skipped(real_translator, fixture_path, tmp_path):
    output_path = str(tmp_path / 'output.docx')
    translate_word_document(
        input_docx_file=fixture_path,
        output_docx_file=output_path,
        source_lang='en',
        use_find_replace=False,
        translation_manager=real_translator
    )
    
    doc = Document(output_path)
    original_doc = Document(fixture_path)
    assert len(doc.paragraphs) == len(original_doc.paragraphs)


@pytest.mark.slow
def test_output_path_generation(real_translator, fixture_path):
    result = translate_word_document(
        input_docx_file=fixture_path,
        output_docx_file=None,
        source_lang='en',
        use_find_replace=False,
        translation_manager=real_translator
    )
    
    assert '_translated_' in result
    assert os.path.exists(result)
    os.remove(result)
    
    notes_path = os.path.splitext(result)[0] + '_translation_notes.json'
    if os.path.exists(notes_path):
        os.remove(notes_path)


@pytest.mark.slow
def test_spacing_between_runs(real_translator, fixture_path, tmp_path):
    output_path = str(tmp_path / 'output.docx')
    translate_word_document(
        input_docx_file=fixture_path,
        output_docx_file=output_path,
        source_lang='en',
        use_find_replace=False,
        translation_manager=real_translator
    )
    
    doc = Document(output_path)
    for paragraph in doc.paragraphs:
        if len(paragraph.runs) <= 1:
            continue
        text = paragraph.text
        if '[NOVALIDTRANSLATIONS]' in text or '[TRANSLATION FAILED]' in text:
            continue
        for word in text.split():
            assert len(word) <= 25, (
                f"Suspicious long word '{word[:50]}' in: {text[:100]}"
            )


# ---------------------------------------------------------------------------
# Unit tests (mock translator, fast)
# ---------------------------------------------------------------------------

def test_long_paragraph_chunking():
    mock = MockTranslator()
    long_text = "This is a test sentence. " * 30  # ~750 chars
    
    doc = Document()
    para = doc.add_paragraph()
    para.add_run(long_text.strip())
    original_length = len(long_text.strip())
    
    _translate_paragraph(para, mock, "en", "fr", False, 1)
    
    assert len(para.text) >= original_length, (
        f"Long paragraph truncated: {len(para.text)} < {original_length}"
    )


@pytest.mark.parametrize("build_fn, expected_link_text", [
    (
            lambda doc: (
                    doc.add_paragraph(),
                    _add_hyperlink(doc.paragraphs[-1], 'https://example.com', 'Our Website'),
                    doc.paragraphs[-1].add_run(' has more details about the project.'),
            ),
            'Our Website',
    ),
    (
            lambda doc: (
                    doc.add_paragraph(),
                    doc.paragraphs[-1].add_run('Publications will be posted on '),
                    _add_hyperlink(doc.paragraphs[-1], 'https://dfo-mpo.gc.ca', 'DFO Science'),
                    doc.paragraphs[-1].add_run(' as they become available.'),
            ),
            'DFO Science',
    ),
    (
            lambda doc: (
                    doc.add_paragraph(),
                    doc.paragraphs[-1].add_run('For more information visit '),
                    _add_hyperlink(doc.paragraphs[-1], 'https://example.com/info', 'the official page'),
            ),
            'the official page',
    ),
], ids=['link_start', 'link_middle', 'link_end'])
def test_hyperlink_in_paragraph(build_fn, expected_link_text, tmp_path):
    mock = MockTranslator()
    input_path = str(tmp_path / 'input.docx')
    output_path = str(tmp_path / 'output.docx')
    
    # Arrange
    doc = Document()
    build_fn(doc)
    doc.save(input_path)
    
    # Act
    translate_word_document(
        input_docx_file=input_path,
        output_docx_file=output_path,
        source_lang="en",
        use_find_replace=False,
        translation_manager=mock
    )
    
    # Assert
    output_doc = Document(output_path)
    output_para = output_doc.paragraphs[0]
    output_text = _get_full_paragraph_text(output_para)
    
    assert expected_link_text in output_text, (
        f"Hyperlink text '{expected_link_text}' missing from output: {output_text}"
    )
    assert '[TR:' in output_text, f"No translated content in output: {output_text}"
    
    # Hyperlink text should not be orphaned after punctuation
    orphan_pattern = re.compile(r'[.!?]' + re.escape(expected_link_text) + r'$')
    assert not orphan_pattern.search(output_text.rstrip()), (
        f"Hyperlink text orphaned after punctuation: {output_text}"
    )
    
    # Hyperlink XML should be stripped
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    assert not output_para._element.findall('.//w:hyperlink', nsmap), (
        "Hyperlink XML was not stripped after translation"
    )


@pytest.mark.parametrize("input_text, expected", [
    ('Figure 1. Map showing the area.', ['Figure 1. Map showing the area.']),
    ('Fig. 3 shows results. The data is clear.', ['Fig. 3 shows results.', 'The data is clear.']),
    ('Tableau 2. Les résultats montrent que...', ['Tableau 2. Les résultats montrent que...']),
    ('Table 1. First sentence. Second sentence.', ['Table 1. First sentence.', 'Second sentence.']),
    ('Hello world. Goodbye world.', ['Hello world.', 'Goodbye world.']),
    ('No periods here', ['No periods here']),
], ids=[
    'figure_label_not_split',
    'fig_abbreviation_protected',
    'french_tableau_not_split',
    'table_label_protected',
    'normal_splitting',
    'no_split_points',
])
def test_split_by_sentences(input_text, expected):
    result, _ = split_by_sentences(input_text)
    assert result == expected


def test_hyperlink_stripping_and_records(tmp_path):
    mock = MockTranslator()
    
    # --- Paragraph WITH hyperlink ---
    temp_path = str(tmp_path / 'hyperlink.docx')
    doc = Document()
    para = doc.add_paragraph()
    para.add_run('Visit ')
    _add_hyperlink(para, 'https://example.com', 'our site')
    para.add_run(' for details.')
    doc.save(temp_path)
    
    doc = Document(temp_path)
    para = doc.paragraphs[0]
    
    formatting_records = []
    _translate_paragraph(para, mock, 'en', 'fr', False, 1, formatting_records=formatting_records)
    
    # No w:hyperlink elements should remain
    assert len(para._element.findall(ns.qn('w:hyperlink'))) == 0, \
        "Hyperlink XML not stripped from paragraph"
    
    # Should have a hyperlink record
    url_records = [r for r in formatting_records if 'https://' in r.get('notes', '')]
    assert len(url_records) >= 1
    record = url_records[0]
    assert record['original_text'] == 'our site'
    assert record['notes'] == 'https://example.com'
    assert 'Visit' in record['full_paragraph'] and 'our site' in record['full_paragraph']
    
    # --- Paragraph WITHOUT hyperlink should NOT get cyan highlighting ---
    doc2 = Document()
    para2 = doc2.add_paragraph()
    para2.add_run('Plain text without any links.')
    
    _translate_paragraph(para2, mock, 'en', 'fr', False, 1, formatting_records=[])
    
    any_cyan = any(
        run.font.highlight_color == WD_COLOR_INDEX.TURQUOISE
        for run in para2.runs
        if hasattr(run, 'font') and hasattr(run.font, 'highlight_color')
    )
    assert not any_cyan, "Cyan highlighting incorrectly applied to paragraph without hyperlinks"


def test_write_formatting_notes(tmp_path):
    records = [
        {'original_text': 'Example Link', 'full_paragraph': 'Visit Example Link for details.', 'notes': 'https://example.com', 'type': 'url'},
        {'original_text': 'Another', 'full_paragraph': 'See Another for more.', 'notes': 'https://another.com', 'type': 'url'},
    ]
    out = str(tmp_path / 'notes.json')
    write_notes_json(records, out)
    
    with open(out, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    total_entries = len(data.get('paragraphs', [])) + len(data.get('tables', []))
    assert total_entries == 2, f"Expected 2 entries, got {total_entries}"
    
    # Check first entry
    first = data['paragraphs'][0]
    assert 'Visit Example Link for details.' == first['full_paragraph']
    assert any('Example Link' in n['original_text'] for n in first['notes'])
    assert any('https://example.com' in n['detail'] for n in first['notes'])


@pytest.mark.parametrize("source_lang, expected_locale", [
    ('en', 'fr-CA'),
    ('fr', 'en-CA'),
], ids=['en_to_fr', 'fr_to_en'])
def test_proofing_language_set(source_lang, expected_locale, tmp_path):
    mock = MockTranslator()
    input_path = str(tmp_path / 'input.docx')
    output_path = str(tmp_path / 'output.docx')
    
    doc = Document()
    doc.add_paragraph("This is a test sentence.")
    doc.save(input_path)
    
    translate_word_document(
        input_docx_file=input_path,
        output_docx_file=output_path,
        source_lang=source_lang,
        use_find_replace=False,
        translation_manager=mock
    )
    
    output_doc = Document(output_path)
    all_r_elements = list(output_doc.element.iter(ns.qn('w:r')))
    assert all_r_elements, "No w:r elements found in output"
    
    for r_elem in all_r_elements:
        rPr = r_elem.find(ns.qn('w:rPr'))
        assert rPr is not None, "w:r element missing rPr"
        lang = rPr.find(ns.qn('w:lang'))
        assert lang is not None, "rPr missing w:lang element"
        val = lang.get(ns.qn('w:val'))
        assert val == expected_locale, f"Expected '{expected_locale}', got '{val}'"
