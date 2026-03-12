import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from scitrans.translate.word_document import _translate_paragraph, _collapse_runs_preserving_shapes
from scitrans.translate.word_formatting import SuperscriptOrdinalsRule, ItalicBracketsRule, apply_formatting_rules
from scitrans.translate.word_notes import _group_notes_by_paragraph
from tests.conftest import run_word_translation, all_notes_text, notes_entry_count

FIXTURE = 'test_formatting_errors_en.docx'


@pytest.fixture(scope="module")
def translated():
    return run_word_translation(FIXTURE, 'en')


@pytest.fixture
def empty_doc():
    return Document()


# ---------------------------------------------------------------------------
# Header / footer structure
# ---------------------------------------------------------------------------

def test_first_page_header_has_tabs(translated):
    doc, _, _ = translated
    first_page_header = doc.sections[0].first_page_header
    tab_found = any(
        run_elem.find(qn('w:tab')) is not None
        for para in first_page_header.paragraphs
        for run_elem in para._element.findall(qn('w:r'))
    )
    assert tab_found, "First page header should contain tab characters in runs"


def test_other_header_has_left_and_right_text(translated):
    doc, _, _ = translated
    header = doc.sections[0].header
    
    assert header.tables, "Non-first-page header should contain a table"
    row = header.tables[0].rows[0]
    assert len(row.cells) >= 2, "Header table should have at least 2 columns"
    
    left_text = row.cells[0].text.strip()
    right_text = row.cells[1].text.strip()
    assert left_text and '[TR:' in left_text, f"Left cell should be translated, got: {left_text}"
    assert right_text and '[TR:' in right_text, f"Right cell should be translated, got: {right_text}"


def test_first_page_footer_image_tabbed_right(translated):
    doc, _, _ = translated
    first_page_footer = doc.sections[0].first_page_footer
    
    has_tab = any(
        para._element.findall('.//' + qn('w:tab'))
        for para in first_page_footer.paragraphs
    )
    has_drawing = any(
        para._element.findall('.//' + qn('w:drawing'))
        for para in first_page_footer.paragraphs
    )
    assert has_tab, "First page footer should have a tab separating text from image"
    assert has_drawing, "First page footer should contain an image/drawing"


def test_first_page_footer_no_page_number(translated):
    doc, _, _ = translated
    first_page_footer = doc.sections[0].first_page_footer
    
    for para in first_page_footer.paragraphs:
        for instr in para._element.findall('.//' + qn('w:instrText')):
            assert 'PAGE' not in (instr.text or '').upper(), \
                "First page footer should not contain a PAGE field code"


def test_page_numbering_is_auto_updating(translated):
    doc, _, _ = translated
    footer = doc.sections[0].footer
    
    has_page_field = any(
        'PAGE' in (instr.text or '').upper()
        for para in footer.paragraphs
        for instr in para._element.findall('.//' + qn('w:instrText'))
    )
    assert has_page_field, (
        "Footer should contain a PAGE field code for auto-updating page numbers"
    )


def test_first_page_footer_preserves_tab_alignment(translated):
    doc, _, _ = translated
    first_page_footer = doc.sections[0].first_page_footer
    
    target_para = next((p for p in first_page_footer.paragraphs if p._element.findall('.//' + qn('w:drawing'))), None)
    assert target_para is not None, "Could not find a paragraph with a drawing in the footer"
    
    elements_in_order = []
    for run in target_para._element.findall(qn('w:r')):
        for child in run:
            if child.tag == qn('w:t') and child.text and child.text.strip():
                elements_in_order.append('text')
            elif child.tag == qn('w:tab'):
                elements_in_order.append('tab')
            elif child.tag == qn('w:drawing'):
                elements_in_order.append('drawing')
    
    assert 'drawing' in elements_in_order, "Drawing not found in runs"
    drawing_idx = elements_in_order.index('drawing')
    
    assert drawing_idx > 0, "Drawing is the first element, expected text and tab before it"
    assert elements_in_order[drawing_idx - 1] == 'tab', (
        f"Expected a tab immediately before the image. "
        f"Full sequence: {elements_in_order}"
    )


def test_collapse_runs_preserves_tabs(empty_doc, subtests):
    with subtests.test('text run, tab run, text run'):
        para = empty_doc.add_paragraph()
        
        para.add_run("Left")
        para.add_run("\t")
        para.add_run("Right")
        
        _collapse_runs_preserving_shapes(para)
        
        assert len(para.runs) == 3, f"Expected 3 isolated runs, got {len(para.runs)}"
        assert len(para.runs[1]._element.findall(qn('w:tab'))) == 1
    
    with subtests.test('text run, tab + text run'):
        para = empty_doc.add_paragraph()
        
        para.add_run("Left")
        para.add_run("\tRight")
        
        _collapse_runs_preserving_shapes(para)
        
        assert len(para.runs[1]._element.findall(qn('w:tab'))) == 1


def test_translate_paragraph_preserves_inline_tabs(empty_doc, mock_translator):
    # Simulate the paragraph state after _collapse_runs_preserving_shapes
    para = empty_doc.add_paragraph()
    para.add_run("Left text")
    
    tab_run = para.add_run()
    tab_run._element.append(OxmlElement('w:tab'))
    
    para.add_run("Right text")
    
    _translate_paragraph(
        paragraph=para,
        translation_manager=mock_translator,
        source_lang="en",
        target_lang="fr",
        use_find_replace=False,
        idx=1,
        use_cache=False,
        formatting_records=[],
        preferential_dict=None,
        chunk_by="sentences",
        location=None
    )
    
    # Verify the tab run was skipped and preserved, while text runs were translated
    assert len(para.runs) == 3
    assert len(para.runs[1]._element.findall(qn('w:tab'))) == 1
    assert "TR:" in para.runs[0].text
    assert "TR:" in para.runs[2].text


# ---------------------------------------------------------------------------
# Superscript / formatting note tests
# ---------------------------------------------------------------------------

def test_superscript_footnote_no_formatting_note(translated):
    _, _, notes_data = translated
    if notes_data is None:
        return
    text = all_notes_text(notes_data)
    assert 'footnote' not in text.lower(), \
        "Superscript in 'footnote1' should not produce a formatting note"


def test_superscript_ordinals_no_formatting_note(translated):
    _, _, notes_data = translated
    if notes_data is None:
        return
    text = all_notes_text(notes_data)
    assert '50th' not in text and '75th' not in text, \
        "Superscript ordinals '50th' and '75th' should not produce formatting notes"


# ---------------------------------------------------------------------------
# Italic / spp. handling
# ---------------------------------------------------------------------------

def test_table_cell_0_1_no_formatting_note(translated):
    doc, _, notes_data = translated
    cell = doc.tables[0].rows[0].cells[1]
    
    has_italic_in_brackets = False
    for para in cell.paragraphs:
        in_bracket = False
        for run in para.runs:
            if '(' in run.text:
                in_bracket = True
            if in_bracket and run.italic and run.text.strip():
                has_italic_in_brackets = True
            if ')' in run.text:
                in_bracket = False
    assert has_italic_in_brackets, (
        "Table cell [0,1] should have italic text inside brackets for species names"
    )
    
    is_highlighted = any(
        run.font.highlight_color is not None
        for para in cell.paragraphs
        for run in para.runs
        if run.text and run.text.strip()
    )
    if is_highlighted and notes_data:
        cell_prefix = 'Shrimp remains an important forage species'
        has_row = any(
            cell_prefix in entry['full_paragraph']
            for section in ['paragraphs', 'tables']
            for entry in notes_data.get(section, [])
        )
        assert has_row, "Cell [0,1] is highlighted but has no corresponding row in notes"


def test_italic_mismatch_note_uses_check_formatting_key(translated):
    _, _, notes_data = translated
    assert notes_data is not None, "Expected notes to be generated"
    
    para_9_prefix = 'This bracket mismatch error'
    found = False
    for section in ['paragraphs', 'tables']:
        for entry in notes_data.get(section, []):
            if para_9_prefix in entry['full_paragraph']:
                found = True
                assert any(
                    n['original_text'] == 'check formatting' for n in entry['notes']
                ), "Italic mismatch note should use 'check formatting' as original_text"
                break
    assert found, "Expected a note for the paragraph with italic mismatch"


def test_final_paragraph_page2_no_formatting_note(translated):
    doc, _, _ = translated
    
    target_para = None
    for para in doc.paragraphs:
        if para.text and 'Shrimp remains' in para.text:
            target_para = para
    
    assert target_para is not None
    is_highlighted = any(
        run.font.highlight_color is not None
        for run in target_para.runs
        if run.text and run.text.strip()
    )
    assert not is_highlighted, (
        "Final paragraph should not be highlighted — italic+spp should be auto-handled"
    )


# ---------------------------------------------------------------------------
# Additional structural tests
# ---------------------------------------------------------------------------

def test_page_breaks_preserved(translated):
    doc, _, _ = translated
    page_breaks = [
        br for para in doc.paragraphs
        for br in para._element.findall('.//' + qn('w:br'))
        if br.get(qn('w:type')) == 'page'
    ]
    assert len(page_breaks) >= 2, (
        f"Expected at least 2 page breaks, found {len(page_breaks)}"
    )


def test_superscript_ordinals_applied_correctly(translated):
    doc, _, _ = translated
    target_para = None
    for para in doc.paragraphs:
        if '50' in para.text and '75' in para.text:
            target_para = para
            break
    assert target_para is not None, "Could not find paragraph with 50th/75th"
    
    runs = target_para.runs
    found_50_super = False
    found_75_super = False
    for i, run in enumerate(runs):
        if run.text.rstrip().endswith('50') and i + 1 < len(runs):
            next_run = runs[i + 1]
            if next_run.font.superscript and next_run.text.startswith('th'):
                found_50_super = True
        if run.text.rstrip().endswith('75') and i + 1 < len(runs):
            next_run = runs[i + 1]
            if next_run.font.superscript and next_run.text.startswith('th'):
                found_75_super = True
    
    assert found_50_super, "50th should have superscript 'th' suffix"
    assert found_75_super, "75th should have superscript 'th' suffix"


def test_highlighted_paragraphs_match_notes_rows(translated):
    doc, _, notes_data = translated
    
    highlighted_count = 0
    seen_texts = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.highlight_color is not None and para.text not in seen_texts:
                highlighted_count += 1
                seen_texts.add(para.text)
                break
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.highlight_color is not None and para.text not in seen_texts:
                            highlighted_count += 1
                            seen_texts.add(para.text)
                            break
    
    assert notes_data is not None, "Expected notes to be generated"
    notes_rows = notes_entry_count(notes_data)
    
    assert highlighted_count == notes_rows, (
        f"Highlighted paragraphs ({highlighted_count}) should equal "
        f"notes entries ({notes_rows})"
    )


# ---------------------------------------------------------------------------
# Superscript ordinals: French "e" suffix must only superscript after the
# target number, not the first "e" found in the run.
# ---------------------------------------------------------------------------

def test_superscript_ordinals_french_targets_correct_position():
    doc = Document()
    para = doc.add_paragraph(
        "Cette note de bas de page ne doit pas entraîner. Les 50e ou 75e percentiles."
    )
    SuperscriptOrdinalsRule().apply(para, ['50th', '75th'])
    
    full_text = ""
    superscripted_positions = []
    for run in para.runs:
        if run.font.superscript:
            for ch in run.text:
                superscripted_positions.append(len(full_text))
                full_text += ch
        else:
            full_text += run.text
    
    idx_50e = full_text.find("50e")
    idx_75e = full_text.find("75e")
    assert idx_50e != -1, f"Could not find '50e' in: {full_text!r}"
    assert idx_75e != -1, f"Could not find '75e' in: {full_text!r}"
    
    expected_positions = {idx_50e + 2, idx_75e + 2}
    actual_positions = set(superscripted_positions)
    
    assert actual_positions == expected_positions, (
        f"Superscript should only be at positions {expected_positions} "
        f"(the 'e' in '50e' and '75e'), but was at {actual_positions}. "
        f"Text: {full_text!r}"
    )


# ---------------------------------------------------------------------------
# Deduplication tests for _group_notes_by_paragraph
# ---------------------------------------------------------------------------

class TestGroupNotesByParagraph:
    def test_same_url_fragments_collapse(self):
        records = [
            {'original_text': 'A', 'full_paragraph': 'Miller full text', 'notes': 'https://example.com/doc', 'type': 'url'},
            {'original_text': 'Near', 'full_paragraph': 'Miller full text', 'notes': 'https://example.com/doc', 'type': 'url'},
            {'original_text': 'l', 'full_paragraph': 'Miller full text', 'notes': 'https://example.com/doc', 'type': 'url'},
            {'original_text': 'y Successful', 'full_paragraph': 'Miller full text', 'notes': 'https://example.com/doc', 'type': 'url'},
        ]
        groups = _group_notes_by_paragraph(records)
        total_rows = sum(len(v) for v in groups.values())
        assert total_rows == 1, f"Same paragraph + same URL should produce 1 row, got {total_rows}"
    
    def test_different_errors_kept(self):
        records = [
            {'original_text': 'bold text', 'full_paragraph': 'Full paragraph', 'notes': 'bold', 'type': 'formatting'},
            {'original_text': 'italic text', 'full_paragraph': 'Full paragraph', 'notes': 'italic', 'type': 'formatting'},
        ]
        groups = _group_notes_by_paragraph(records)
        total_rows = sum(len(v) for v in groups.values())
        assert total_rows == 2, f"Different errors should produce 2 rows, got {total_rows}"
    
    def test_different_paragraphs_not_collapsed(self):
        records = [
            {'original_text': 'text', 'full_paragraph': 'First paragraph text', 'notes': 'italic', 'type': 'formatting'},
            {'original_text': 'text', 'full_paragraph': 'Second paragraph text', 'notes': 'italic', 'type': 'formatting'},
        ]
        groups = _group_notes_by_paragraph(records)
        total_rows = sum(len(v) for v in groups.values())
        assert total_rows == 2, f"Different paragraphs should each get a row, got {total_rows}"
    
    def test_exact_duplicates_collapsed(self):
        records = [
            {'original_text': 'text', 'full_paragraph': 'Same paragraph', 'notes': 'italic', 'type': 'formatting'},
            {'original_text': 'text', 'full_paragraph': 'Same paragraph', 'notes': 'italic', 'type': 'formatting'},
            {'original_text': 'text', 'full_paragraph': 'Same paragraph', 'notes': 'italic', 'type': 'formatting'},
        ]
        groups = _group_notes_by_paragraph(records)
        total_rows = sum(len(v) for v in groups.values())
        assert total_rows == 1, f"Exact duplicates should collapse to 1, got {total_rows}"
    
    def test_same_notes_no_location_collapsed(self):
        records = [
            {'original_text': 'text', 'full_paragraph': 'Same paragraph', 'notes': 'italic', 'type': 'formatting'},
            {'original_text': 'text', 'full_paragraph': 'Same paragraph', 'notes': 'italic', 'type': 'formatting'},
        ]
        groups = _group_notes_by_paragraph(records)
        total_rows = sum(len(v) for v in groups.values())
        assert total_rows == 1, f"Same notes without location should collapse to 1, got {total_rows}"
    
    def test_same_text_different_locations_not_collapsed(self):
        # Same paragraph text in body and table cell should get separate notes entries
        records = [
            {
                'original_text': 'check formatting',
                'full_paragraph': 'Shrimp remains an important forage species...',
                'notes': 'Italic bracket formatting could not be automatically applied',
                'type': 'formatting',
                'location': {'section': 'paragraphs', 'index': 11},
            },
            {
                'original_text': 'check formatting',
                'full_paragraph': 'Shrimp remains an important forage species...',
                'notes': 'Italic bracket formatting could not be automatically applied',
                'type': 'formatting',
                'location': {'section': 'tables', 'table': 0, 'row': 0, 'cell': 1},
            },
        ]
        groups = _group_notes_by_paragraph(records)
        total_rows = sum(len(v) for v in groups.values())
        assert total_rows == 2, (
            f"Same text in different locations (body vs table) should produce "
            f"2 separate entries, got {total_rows}"
        )


def test_italic_brackets_with_extra_bracket_from_translation():
    # Real translators sometimes add extra bracketed text (e.g. "(less formal)")
    # that wasn't in the source. The italic bracket rule should still apply
    # italic to the species-name brackets.
    doc = Document()
    translated_text = (
        "La crevette demeure dans le NL bioregion "
        "(less formal). Les prédateurs comprennent Greenland Halibut "
        "(Reinhardtius hippoglossoides), Atlantic Cod (Gadus morhua) et "
        "Redfish (Sebastes spp.)."
    )
    para = doc.add_paragraph(translated_text)
    
    # Source had 3 italic brackets — pass detected patterns from source
    italic_rule = ItalicBracketsRule()
    detected = {italic_rule: 3}
    
    formatting_records = []
    apply_formatting_rules(para, formatting_records, "source text", detected=detected)
    
    # The species name brackets should have italic content
    has_italic = any(run.italic for run in para.runs if run.text.strip())
    assert has_italic, (
        "Italic bracket rule should still apply italic to species brackets "
        "even when translation adds extra non-species brackets"
    )
    
    # Should NOT have a "check formatting" note for species brackets
    check_notes = [
        r for r in formatting_records
        if 'check formatting' in r.get('original_text', '')
    ]
    assert len(check_notes) == 0, (
        f"Should not produce 'check formatting' note when species brackets "
        f"are clearly identifiable. Got: {check_notes}"
    )


def test_italic_brackets_fewer_than_expected_produces_note():
    doc = Document()
    para = doc.add_paragraph("Only one bracket (species name) in the text.")
    
    italic_rule = ItalicBracketsRule()
    detected = {italic_rule: 3}
    
    formatting_records = []
    apply_formatting_rules(para, formatting_records, "source text", detected=detected)
    
    check_notes = [
        r for r in formatting_records
        if 'check formatting' in r.get('original_text', '')
    ]
    assert len(check_notes) == 1, (
        f"Should produce 'check formatting' note when fewer brackets than expected. Got: {check_notes}"
    )


def test_italic_brackets_exact_count_applies_correctly():
    doc = Document()
    para = doc.add_paragraph("Species include (Gadus morhua) and (Sebastes spp.).")
    
    italic_rule = ItalicBracketsRule()
    detected = {italic_rule: 2}
    
    formatting_records = []
    apply_formatting_rules(para, formatting_records, "source text", detected=detected)
    
    has_italic = any(run.italic for run in para.runs if run.text.strip())
    assert has_italic, "Italic should be applied when bracket count matches exactly"
    
    check_notes = [r for r in formatting_records if 'check formatting' in r.get('original_text', '')]
    assert len(check_notes) == 0, "No notes should be produced when count matches"
