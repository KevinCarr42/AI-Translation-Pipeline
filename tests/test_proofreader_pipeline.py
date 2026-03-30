import json
from pathlib import Path

import docx
import pytest
from lxml import etree

from scitrans.proofreader.extract_text import extract_text_with_ids, extract_locations
from scitrans.proofreader.accept_changes import (
    accept_all_changes, _accept_insertions, _accept_deletions,
    W_NS, W,
)
from scitrans.proofreader.apply_review import apply_tracked_change, build_location_maps
from scitrans.proofreader.lexical_checklist import (
    lexical_constraint_checklist, save_checklist,
)

# Import pipeline helpers from the script
import importlib
import sys

_pipeline_spec = importlib.util.spec_from_file_location(
    'run_proofreader_pipeline',
    Path(__file__).resolve().parent.parent / 'scripts' / 'run_proofreader_pipeline.py',
)
_pipeline_mod = importlib.util.module_from_spec(_pipeline_spec)
sys.modules['run_proofreader_pipeline'] = _pipeline_mod
_pipeline_spec.loader.exec_module(_pipeline_mod)

_make_checkpoint_path = _pipeline_mod._make_checkpoint_path
_load_response_json = _pipeline_mod._load_response_json
STEP_SUFFIXES = _pipeline_mod.STEP_SUFFIXES


# ── Helpers ──────────────────────────────────────────────────────────────

def _create_simple_docx(path, paragraphs=None, table_rows=None):
    doc = docx.Document()
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)
    if table_rows:
        num_cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        for r_idx, row_data in enumerate(table_rows):
            for c_idx, cell_text in enumerate(row_data):
                table.rows[r_idx].cells[c_idx].text = cell_text
    doc.save(str(path))
    return path


def _create_glossary_json(path):
    data = {
        "translations": {
            "nomenclature": [
                {"english": "groundfish", "french": "poisson de fond"},
                {"english": "herring", "french": "hareng"},
            ],
            "acronym": [
                {"english_acronym": "DFO", "french_acronym": "MPO"},
            ],
            "taxon": [],
            "site": [],
            "table": [],
        }
    }
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return path


# ── extract_text_with_ids ────────────────────────────────────────────────

class TestExtractTextWithIds:
    
    def test_paragraphs_and_table(self, tmp_path):
        docx_path = _create_simple_docx(
            tmp_path / 'test.docx',
            paragraphs=['First paragraph', 'Second paragraph'],
            table_rows=[['A1', 'B1'], ['A2', 'B2']],
        )
        
        result = extract_text_with_ids(str(docx_path))
        
        assert '[P0] First paragraph' in result
        assert '[P1] Second paragraph' in result
        assert '[T0-R0] A1 | B1' in result
        assert '[T0-R1] A2 | B2' in result
    
    def test_empty_paragraphs_skipped(self, tmp_path):
        docx_path = _create_simple_docx(
            tmp_path / 'test.docx',
            paragraphs=['Hello', '', '   ', 'World'],
        )
        
        result = extract_text_with_ids(str(docx_path))
        
        assert '[P0] Hello' in result
        assert '[P3] World' in result
        # Empty paragraphs at index 1 and 2 should not appear
        assert '[P1]' not in result
        assert '[P2]' not in result
    
    def test_output_is_newline_joined(self, tmp_path):
        docx_path = _create_simple_docx(
            tmp_path / 'test.docx',
            paragraphs=['Line one', 'Line two'],
        )
        
        result = extract_text_with_ids(str(docx_path))
        lines = result.split('\n')
        
        assert len(lines) == 2


# ── extract_locations ────────────────────────────────────────────────────

class TestExtractLocations:
    
    def test_returns_tuples(self, tmp_path):
        docx_path = _create_simple_docx(
            tmp_path / 'test.docx',
            paragraphs=['Alpha', 'Beta'],
            table_rows=[['C1', 'C2']],
        )
        
        locations = extract_locations(str(docx_path))
        
        assert ('P0', 'Alpha') in locations
        assert ('P1', 'Beta') in locations
        assert ('T0-R0', 'C1 | C2') in locations
    
    def test_skips_empty_paragraphs(self, tmp_path):
        docx_path = _create_simple_docx(
            tmp_path / 'test.docx',
            paragraphs=['Content', '', 'More content'],
        )
        
        locations = extract_locations(str(docx_path))
        loc_ids = [loc_id for loc_id, _ in locations]
        
        assert 'P0' in loc_ids
        assert 'P2' in loc_ids
        assert 'P1' not in loc_ids


# ── accept_changes ───────────────────────────────────────────────────────

class TestAcceptChanges:
    
    def _inject_tracked_insertion(self, doc, paragraph_text, inserted_text):
        # Build a paragraph with a tracked insertion using raw XML
        p = doc.add_paragraph()
        # Normal run with paragraph_text
        run_xml = (
            f'<w:r xmlns:w="{W_NS}">'
            f'<w:t xml:space="preserve">{paragraph_text}</w:t>'
            f'</w:r>'
        )
        p._element.append(etree.fromstring(run_xml))
        # Tracked insertion
        ins_xml = (
            f'<w:ins xmlns:w="{W_NS}" w:author="test" w:date="2026-01-01T00:00:00Z">'
            f'<w:r><w:t xml:space="preserve">{inserted_text}</w:t></w:r>'
            f'</w:ins>'
        )
        p._element.append(etree.fromstring(ins_xml))
        return p
    
    def _inject_tracked_deletion(self, doc, visible_text, deleted_text):
        p = doc.add_paragraph()
        run_xml = (
            f'<w:r xmlns:w="{W_NS}">'
            f'<w:t xml:space="preserve">{visible_text}</w:t>'
            f'</w:r>'
        )
        p._element.append(etree.fromstring(run_xml))
        del_xml = (
            f'<w:del xmlns:w="{W_NS}" w:author="test" w:date="2026-01-01T00:00:00Z">'
            f'<w:r><w:delText xml:space="preserve">{deleted_text}</w:delText></w:r>'
            f'</w:del>'
        )
        p._element.append(etree.fromstring(del_xml))
        return p
    
    def test_accept_insertions(self, tmp_path):
        doc = docx.Document()
        self._inject_tracked_insertion(doc, 'Hello ', 'World')
        input_path = tmp_path / 'input.docx'
        output_path = tmp_path / 'output.docx'
        doc.save(str(input_path))
        
        result = accept_all_changes(str(input_path), str(output_path))
        
        assert result['insertions_accepted'] == 1
        accepted_doc = docx.Document(str(output_path))
        # The inserted text should be kept (unwrapped from w:ins)
        full_text = ''.join(p.text for p in accepted_doc.paragraphs)
        assert 'World' in full_text
    
    def test_accept_deletions(self, tmp_path):
        doc = docx.Document()
        self._inject_tracked_deletion(doc, 'Keep this', ' Remove this')
        input_path = tmp_path / 'input.docx'
        output_path = tmp_path / 'output.docx'
        doc.save(str(input_path))
        
        result = accept_all_changes(str(input_path), str(output_path))
        
        assert result['deletions_accepted'] == 1
        accepted_doc = docx.Document(str(output_path))
        full_text = ''.join(p.text for p in accepted_doc.paragraphs)
        assert 'Keep this' in full_text
        assert 'Remove this' not in full_text
    
    def test_accept_all_returns_zero_on_clean_doc(self, tmp_path):
        doc = docx.Document()
        doc.add_paragraph('No tracked changes here')
        input_path = tmp_path / 'clean.docx'
        output_path = tmp_path / 'output.docx'
        doc.save(str(input_path))
        
        result = accept_all_changes(str(input_path), str(output_path))
        
        assert result['insertions_accepted'] == 0
        assert result['deletions_accepted'] == 0
        assert result['format_changes_accepted'] == 0
        assert result['paragraph_changes_accepted'] == 0
    
    def test_accept_format_changes(self, tmp_path):
        doc = docx.Document()
        p = doc.add_paragraph()
        # Run with a formatting change tracked via rPrChange
        run_xml = (
            f'<w:r xmlns:w="{W_NS}">'
            f'<w:rPr><w:b/>'
            f'<w:rPrChange w:author="test" w:date="2026-01-01T00:00:00Z">'
            f'<w:rPr/>'
            f'</w:rPrChange>'
            f'</w:rPr>'
            f'<w:t>Bold text</w:t>'
            f'</w:r>'
        )
        p._element.append(etree.fromstring(run_xml))
        input_path = tmp_path / 'fmt.docx'
        output_path = tmp_path / 'output.docx'
        doc.save(str(input_path))
        
        result = accept_all_changes(str(input_path), str(output_path))
        
        assert result['format_changes_accepted'] == 1


# ── lexical_constraint_checklist ─────────────────────────────────────────

class TestLexicalChecklist:
    
    def test_finds_glossary_terms(self, tmp_path):
        glossary_path = _create_glossary_json(tmp_path / 'glossary.json')
        docx_path = _create_simple_docx(
            tmp_path / 'test.docx',
            paragraphs=[
                'The groundfish survey was conducted by DFO.',
                'Herring populations are declining.',
                'No matching terms here.',
            ],
        )
        
        checklist = lexical_constraint_checklist(
            str(docx_path), source_lang='en', glossary_path=str(glossary_path),
        )
        
        locations = {(item['location'], item['source_text']) for item in checklist}
        assert ('P0', 'groundfish') in locations
        assert ('P0', 'DFO') in locations
        assert ('P1', 'herring') in locations
    
    def test_no_matches_returns_empty(self, tmp_path):
        glossary_path = _create_glossary_json(tmp_path / 'glossary.json')
        docx_path = _create_simple_docx(
            tmp_path / 'test.docx',
            paragraphs=['Nothing relevant here at all.'],
        )
        
        checklist = lexical_constraint_checklist(
            str(docx_path), source_lang='en', glossary_path=str(glossary_path),
        )
        
        assert checklist == []
    
    def test_case_insensitive_matching(self, tmp_path):
        glossary_path = _create_glossary_json(tmp_path / 'glossary.json')
        docx_path = _create_simple_docx(
            tmp_path / 'test.docx',
            paragraphs=['GROUNDFISH are important.'],
        )
        
        checklist = lexical_constraint_checklist(
            str(docx_path), source_lang='en', glossary_path=str(glossary_path),
        )
        
        assert len(checklist) == 1
        assert checklist[0]['source_text'] == 'groundfish'
    
    def test_finds_terms_in_tables(self, tmp_path):
        glossary_path = _create_glossary_json(tmp_path / 'glossary.json')
        docx_path = _create_simple_docx(
            tmp_path / 'test.docx',
            paragraphs=['Intro text'],
            table_rows=[['herring count', '500']],
        )
        
        checklist = lexical_constraint_checklist(
            str(docx_path), source_lang='en', glossary_path=str(glossary_path),
        )
        
        table_matches = [item for item in checklist if item['location'].startswith('T')]
        assert len(table_matches) == 1
        assert table_matches[0]['source_text'] == 'herring'


# ── save_checklist ───────────────────────────────────────────────────────

class TestSaveChecklist:
    
    def test_round_trip(self, tmp_path):
        checklist = [
            {"location": "P0", "source_text": "groundfish", "preferred_translation": "poisson de fond"},
            {"location": "P1", "source_text": "DFO", "preferred_translation": "MPO"},
        ]
        output_path = tmp_path / 'checklist.json'
        
        result_path = save_checklist(checklist, str(output_path))
        
        assert result_path == str(output_path)
        with open(output_path, 'r', encoding='utf-8') as f:
            loaded = json.load(f)
        assert loaded == checklist
    
    def test_unicode_preserved(self, tmp_path):
        checklist = [
            {"location": "P0", "source_text": "poisson", "preferred_translation": "poisson de fond"},
        ]
        output_path = tmp_path / 'checklist.json'
        
        save_checklist(checklist, str(output_path))
        
        raw = output_path.read_text(encoding='utf-8')
        # ensure_ascii=False means accented chars are preserved literally
        assert 'poisson de fond' in raw


# ── _make_checkpoint_path ────────────────────────────────────────────────

class TestMakeCheckpointPath:
    
    @pytest.mark.parametrize("input_name, step_suffix, expected_name", [
        ('1432_en_translated.docx', '_fix_formatting', '1432_en_translated_fix_formatting.docx'),
        ('1432_en_translated_fix_formatting.docx', '_proofreading', '1432_en_translated_proofreading.docx'),
        ('1432_en_translated_proofreading.docx', '_lexical_constraints', '1432_en_translated_lexical_constraints.docx'),
        ('1432_en_translated_lexical_constraints.docx', '_recommended_updates', '1432_en_translated_recommended_updates.docx'),
        ('1432_en_translated_recommended_updates.docx', '_fix_formatting', '1432_en_translated_fix_formatting.docx'),
    ])
    def test_checkpoint_path(self, tmp_path, input_name, step_suffix, expected_name):
        input_path = tmp_path / input_name
        
        result = _make_checkpoint_path(input_path, step_suffix)
        
        assert result == tmp_path / expected_name
    
    def test_no_existing_suffix_appends(self, tmp_path):
        input_path = tmp_path / 'report_translated.docx'
        
        result = _make_checkpoint_path(input_path, '_fix_formatting')
        
        assert result.name == 'report_translated_fix_formatting.docx'


# ── _load_response_json ──────────────────────────────────────────────────

class TestLoadResponseJson:
    
    def test_raw_json(self, tmp_path):
        data = [{"location": "P0", "original": "old", "replacement": "new"}]
        response_path = tmp_path / 'response.json'
        response_path.write_text(json.dumps(data), encoding='utf-8')
        
        result = _load_response_json(str(response_path))
        
        assert result == data
    
    def test_fenced_json(self, tmp_path):
        data = [{"location": "P0", "original": "old", "replacement": "new"}]
        fenced = f'```json\n{json.dumps(data, indent=2)}\n```'
        response_path = tmp_path / 'response.json'
        response_path.write_text(fenced, encoding='utf-8')
        
        result = _load_response_json(str(response_path))
        
        assert result == data
    
    def test_fenced_without_language_tag(self, tmp_path):
        data = {"key": "value"}
        fenced = f'```\n{json.dumps(data)}\n```'
        response_path = tmp_path / 'response.json'
        response_path.write_text(fenced, encoding='utf-8')
        
        result = _load_response_json(str(response_path))
        
        assert result == data
    
    def test_whitespace_around_fences(self, tmp_path):
        data = [1, 2, 3]
        fenced = f'  ```json\n{json.dumps(data)}\n```  '
        response_path = tmp_path / 'response.json'
        response_path.write_text(fenced, encoding='utf-8')

        result = _load_response_json(str(response_path))

        assert result == data


# ── apply_tracked_change ────────────────────────────────────────────────

def _make_para_with_runs(texts):
    p = etree.Element(f'{{{W_NS}}}p')
    for text in texts:
        r = etree.SubElement(p, f'{{{W_NS}}}r')
        t = etree.SubElement(r, f'{{{W_NS}}}t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = text
    return p


def _get_del_text(para):
    for dt in para.iter(f'{{{W_NS}}}delText'):
        return dt.text
    return None


def _get_ins_text(para):
    for t in para.iter(f'{{{W_NS}}}ins'):
        for run_t in t.iter(f'{{{W_NS}}}t'):
            return run_t.text
    return None


DATE = '2026-01-01T00:00:00Z'


class TestApplyTrackedChange:

    def test_exact_match(self):
        p = _make_para_with_runs(['Hello world'])
        result = apply_tracked_change(p, 'world', 'terre', 100, date=DATE)
        assert result is True
        assert _get_del_text(p) == 'world'
        assert _get_ins_text(p) == 'terre'

    def test_nbsp_normalized(self):
        p = _make_para_with_runs(['non\u00a0breaking text'])
        result = apply_tracked_change(p, 'non breaking', 'sans coupure', 100, date=DATE)
        assert result is True
        assert _get_ins_text(p) == 'sans coupure'

    def test_multi_space_normalized(self):
        p = _make_para_with_runs(['double  space here'])
        result = apply_tracked_change(p, 'double space', 'espace double', 100, date=DATE)
        assert result is True
        assert _get_ins_text(p) == 'espace double'

    def test_smart_quotes_normalized(self):
        p = _make_para_with_runs(['\u201cquoted\u201d text'])
        result = apply_tracked_change(p, '"quoted" text', 'texte cit\u00e9', 100, date=DATE)
        assert result is True

    def test_del_text_uses_original_chars(self):
        p = _make_para_with_runs(['a\u00a0b c'])
        result = apply_tracked_change(p, 'a b', 'x y', 100, date=DATE)
        assert result is True
        # delText must contain the original NBSP, not a regular space
        assert _get_del_text(p) == 'a\u00a0b'

    def test_no_match_returns_false(self):
        p = _make_para_with_runs(['Hello world'])
        result = apply_tracked_change(p, 'not here', 'replacement', 100, date=DATE)
        assert result is False

    def test_combined_normalizations(self):
        p = _make_para_with_runs(['\u201chello\u201d\u00a0world'])
        result = apply_tracked_change(p, '"hello" world', 'remplacement', 100, date=DATE)
        assert result is True
        assert _get_del_text(p) == '\u201chello\u201d\u00a0world'


# ── build_location_maps (header/footer support) ────────────────────────

def _create_docx_with_header(path, paragraphs, header_text):
    doc = docx.Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header_para = section.header.paragraphs[0]
    header_para.text = header_text
    doc.save(str(path))
    return path


class TestBuildLocationMaps:

    def test_para_map(self, tmp_path):
        path = _create_docx_with_header(
            tmp_path / 'test.docx', ['Hello world'], 'My Header')
        doc = docx.Document(str(path))
        para_map, table_map, hf_map = build_location_maps(doc)
        assert 'P0' in para_map

    def test_hf_map_has_header(self, tmp_path):
        path = _create_docx_with_header(
            tmp_path / 'test.docx', ['Body text'], 'Header Content')
        doc = docx.Document(str(path))
        para_map, table_map, hf_map = build_location_maps(doc)
        assert 'H0' in hf_map

    def test_hf_map_ids_match_extract_text(self, tmp_path):
        path = _create_docx_with_header(
            tmp_path / 'test.docx', ['Body text'], 'Header Content')
        # Verify that extract_text and build_location_maps assign the same IDs
        from scitrans.proofreader.extract_text import extract_text_with_ids
        extracted = extract_text_with_ids(str(path))
        doc = docx.Document(str(path))
        _, _, hf_map = build_location_maps(doc)
        for loc_id in hf_map:
            assert f'[{loc_id}]' in extracted


# ── fix_formatting header/footer support ────────────────────────────────

class TestFixFormattingHeaderFooter:

    def test_punctuation_rules_applied_to_header(self, tmp_path):
        from scitrans.proofreader.fix_formatting import fix_formatting
        # French rule: space before colon
        path = _create_docx_with_header(
            tmp_path / 'input.docx', ['Body text'], 'Titre: valeur')
        output = tmp_path / 'output.docx'
        fix_formatting(str(path), str(output), lang='fr', use_glossary=False)
        doc = docx.Document(str(output))
        header_text = doc.sections[0].header.paragraphs[0].text
        assert '\u00a0:' in header_text

    def test_glossary_replacements_applied_to_header(self, tmp_path):
        from scitrans.proofreader.fix_formatting import fix_formatting
        glossary_path = _create_glossary_json(tmp_path / 'glossary.json')
        # Create source doc with "DFO" so it appears in sub_glossary
        source_path = _create_simple_docx(
            tmp_path / 'source.docx', paragraphs=['DFO report'])
        # Create translated doc with "DFO" in the header
        path = _create_docx_with_header(
            tmp_path / 'input.docx', ['Rapport'], 'DFO Region')
        output = tmp_path / 'output.docx'
        fix_formatting(
            str(path), str(output), lang='fr',
            source_lang='en', source_path=str(source_path),
            glossary_path=str(glossary_path), use_glossary=True)
        doc = docx.Document(str(output))
        header_text = doc.sections[0].header.paragraphs[0].text
        assert 'MPO' in header_text
