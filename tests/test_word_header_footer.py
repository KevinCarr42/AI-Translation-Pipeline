import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from translate.document import translate_word_document
from docx import Document
import tempfile


class MockTranslator:
    def __init__(self):
        self.call_count = 0
        self.source_texts = []

    def translate_with_best_model(self, text, source_lang, target_lang, use_find_replace, idx, **kwargs):
        self.call_count += 1
        self.source_texts.append(text)
        return {"translated_text": f"[TR:{text}]"}


def _run_translation(fixture_name, source_lang):
    fixture_path = os.path.join(os.path.dirname(__file__), 'fixtures', fixture_name)
    
    temp_output = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    temp_output.close()
    output_path = temp_output.name
    
    mock = MockTranslator()
    
    translate_word_document(
        input_docx_file=fixture_path,
        output_docx_file=output_path,
        source_lang=source_lang,
        use_find_replace=False,
        translation_manager=mock
    )
    
    doc = Document(output_path)
    os.remove(output_path)
    
    return doc, mock


def test_header_paragraph_translated():
    print("\n=== Testing header paragraph translated ===\n")
    
    test_cases = [
        {'name': 'EN fixture - section 2 header translated', 'fixture': 'test_document_structure_en.docx', 'source_lang': 'en'},
        {'name': 'FR fixture - section 2 header translated', 'fixture': 'test_document_structure_fr.docx', 'source_lang': 'fr'},
    ]
    
    passed = 0
    failed = 0
    
    for test in test_cases:
        doc, mock = _run_translation(test['fixture'], test['source_lang'])
        
        section_2_header = doc.sections[2].header
        if section_2_header.paragraphs:
            header_text = section_2_header.paragraphs[0].text
            if '[TR:' in header_text:
                print(f"[PASS] {test['name']}")
                print(f"  Header text: {header_text[:50]}")
                passed += 1
            else:
                print(f"[FAIL] {test['name']}")
                print(f"  Expected: Text containing '[TR:'")
                print(f"  Got: {header_text}")
                failed += 1
        else:
            print(f"[FAIL] {test['name']}")
            print(f"  Expected: Header with paragraphs")
            print(f"  Got: No paragraphs in header")
            failed += 1
    
    print(f"\n{passed} passed, {failed} failed\n")
    assert failed == 0, f"{failed} test cases failed"


def test_footer_paragraphs_translated():
    print("\n=== Testing footer paragraphs translated ===\n")
    
    test_cases = [
        {'name': 'EN fixture - all footers translated', 'fixture': 'test_document_structure_en.docx', 'source_lang': 'en'},
        {'name': 'FR fixture - all footers translated', 'fixture': 'test_document_structure_fr.docx', 'source_lang': 'fr'},
    ]
    
    passed = 0
    failed = 0
    
    for test in test_cases:
        doc, mock = _run_translation(test['fixture'], test['source_lang'])
        
        all_footers_ok = True
        for section_idx in [0, 1, 2]:
            footer = doc.sections[section_idx].footer
            if footer.paragraphs:
                footer_text = footer.paragraphs[0].text
                if '[TR:' not in footer_text:
                    print(f"[FAIL] {test['name']}")
                    print(f"  Section {section_idx} footer not translated: {footer_text}")
                    all_footers_ok = False
                    break
            else:
                print(f"[FAIL] {test['name']}")
                print(f"  Section {section_idx} footer has no paragraphs")
                all_footers_ok = False
                break
        
        if all_footers_ok:
            print(f"[PASS] {test['name']}")
            print(f"  All section footers translated")
            passed += 1
        else:
            failed += 1
    
    print(f"\n{passed} passed, {failed} failed\n")
    assert failed == 0, f"{failed} test cases failed"


def test_header_table_cells_translated():
    print("\n=== Testing header table cells translated ===\n")
    
    test_cases = [
        {'name': 'EN fixture - header table cells translated', 'fixture': 'test_document_structure_en.docx', 'source_lang': 'en'},
        {'name': 'FR fixture - header table cells translated', 'fixture': 'test_document_structure_fr.docx', 'source_lang': 'fr'},
    ]
    
    passed = 0
    failed = 0
    
    for test in test_cases:
        doc, mock = _run_translation(test['fixture'], test['source_lang'])
        
        section_0_header = doc.sections[0].header
        if section_0_header.tables:
            table = section_0_header.tables[0]
            if len(table.rows) > 0 and len(table.rows[0].cells) >= 2:
                cell_0_text = table.rows[0].cells[0].text
                cell_1_text = table.rows[0].cells[1].text
                
                if '[TR:' in cell_0_text and '[TR:' in cell_1_text:
                    print(f"[PASS] {test['name']}")
                    print(f"  Cell 0: {cell_0_text[:40]}")
                    print(f"  Cell 1: {cell_1_text[:40]}")
                    passed += 1
                else:
                    print(f"[FAIL] {test['name']}")
                    print(f"  Expected: Both cells contain '[TR:'")
                    print(f"  Cell 0: {cell_0_text}")
                    print(f"  Cell 1: {cell_1_text}")
                    failed += 1
            else:
                print(f"[FAIL] {test['name']}")
                print(f"  Expected: Table with at least 1 row and 2 cells")
                failed += 1
        else:
            print(f"[FAIL] {test['name']}")
            print(f"  Expected: Header with table")
            failed += 1
    
    print(f"\n{passed} passed, {failed} failed\n")
    assert failed == 0, f"{failed} test cases failed"


def test_first_page_header_translated():
    print("\n=== Testing first page header translated ===\n")
    
    test_cases = [
        {'name': 'EN fixture - first page header translated', 'fixture': 'test_document_structure_en.docx', 'source_lang': 'en'},
        {'name': 'FR fixture - first page header translated', 'fixture': 'test_document_structure_fr.docx', 'source_lang': 'fr'},
    ]
    
    passed = 0
    failed = 0
    
    for test in test_cases:
        doc, mock = _run_translation(test['fixture'], test['source_lang'])
        
        first_page_header = doc.sections[0].first_page_header
        if first_page_header.paragraphs:
            header_text = first_page_header.paragraphs[0].text
            if '[TR:' in header_text:
                print(f"[PASS] {test['name']}")
                print(f"  First page header: {header_text[:50]}")
                passed += 1
            else:
                print(f"[FAIL] {test['name']}")
                print(f"  Expected: Text containing '[TR:'")
                print(f"  Got: {header_text}")
                failed += 1
        else:
            print(f"[FAIL] {test['name']}")
            print(f"  Expected: First page header with paragraphs")
            failed += 1
    
    print(f"\n{passed} passed, {failed} failed\n")
    assert failed == 0, f"{failed} test cases failed"


def test_first_page_footer_translated():
    print("\n=== Testing first page footer translated ===\n")
    
    test_cases = [
        {'name': 'EN fixture - first page footer translated', 'fixture': 'test_document_structure_en.docx', 'source_lang': 'en'},
        {'name': 'FR fixture - first page footer translated', 'fixture': 'test_document_structure_fr.docx', 'source_lang': 'fr'},
    ]
    
    passed = 0
    failed = 0
    
    for test in test_cases:
        doc, mock = _run_translation(test['fixture'], test['source_lang'])
        
        first_page_footer = doc.sections[0].first_page_footer
        if first_page_footer.paragraphs:
            footer_text = first_page_footer.paragraphs[0].text
            if '[TR:' in footer_text:
                print(f"[PASS] {test['name']}")
                print(f"  First page footer: {footer_text[:50]}")
                passed += 1
            else:
                print(f"[FAIL] {test['name']}")
                print(f"  Expected: Text containing '[TR:'")
                print(f"  Got: {footer_text}")
                failed += 1
        else:
            print(f"[FAIL] {test['name']}")
            print(f"  Expected: First page footer with paragraphs")
            failed += 1
    
    print(f"\n{passed} passed, {failed} failed\n")
    assert failed == 0, f"{failed} test cases failed"


def test_even_page_header_translated():
    print("\n=== Testing even page header translated ===\n")
    
    test_cases = [
        {'name': 'EN fixture - even page header translated', 'fixture': 'test_document_structure_en.docx', 'source_lang': 'en'},
        {'name': 'FR fixture - even page header translated', 'fixture': 'test_document_structure_fr.docx', 'source_lang': 'fr'},
    ]
    
    passed = 0
    failed = 0
    
    for test in test_cases:
        doc, mock = _run_translation(test['fixture'], test['source_lang'])
        
        even_page_header = doc.sections[2].even_page_header
        if even_page_header.paragraphs:
            header_text = even_page_header.paragraphs[0].text
            if '[TR:' in header_text:
                print(f"[PASS] {test['name']}")
                print(f"  Even page header: {header_text[:50]}")
                passed += 1
            else:
                print(f"[FAIL] {test['name']}")
                print(f"  Expected: Text containing '[TR:'")
                print(f"  Got: {header_text}")
                failed += 1
        else:
            print(f"[FAIL] {test['name']}")
            print(f"  Expected: Even page header with paragraphs")
            failed += 1
    
    print(f"\n{passed} passed, {failed} failed\n")
    assert failed == 0, f"{failed} test cases failed"


def test_linked_header_not_double_translated():
    print("\n=== Testing linked headers not double-translated ===\n")
    
    test_cases = [
        {'name': 'EN fixture - no double translation', 'fixture': 'test_document_structure_en.docx', 'source_lang': 'en'},
        {'name': 'FR fixture - no double translation', 'fixture': 'test_document_structure_fr.docx', 'source_lang': 'fr'},
    ]
    
    passed = 0
    failed = 0
    
    for test in test_cases:
        doc, mock = _run_translation(test['fixture'], test['source_lang'])
        
        double_translated = [text for text in mock.source_texts if text.startswith('[TR:')]
        
        if len(double_translated) == 0:
            print(f"[PASS] {test['name']}")
            print(f"  No source texts start with '[TR:'")
            passed += 1
        else:
            print(f"[FAIL] {test['name']}")
            print(f"  Expected: No source texts starting with '[TR:'")
            print(f"  Got: {len(double_translated)} texts")
            print(f"  Examples: {double_translated[:3]}")
            failed += 1
    
    print(f"\n{passed} passed, {failed} failed\n")
    assert failed == 0, f"{failed} test cases failed"


def test_linked_header_skipped():
    print("\n=== Testing linked header skipped ===\n")
    
    test_cases = [
        {'name': 'EN fixture - section 1 header is same object as section 0', 'fixture': 'test_document_structure_en.docx', 'source_lang': 'en'},
        {'name': 'FR fixture - section 1 header is same object as section 0', 'fixture': 'test_document_structure_fr.docx', 'source_lang': 'fr'},
    ]
    
    passed = 0
    failed = 0
    
    for test in test_cases:
        doc, mock = _run_translation(test['fixture'], test['source_lang'])
        
        section_0_header = doc.sections[0].header
        section_1_header = doc.sections[1].header
        
        if section_0_header.tables and section_1_header.tables:
            s0_cell_text = section_0_header.tables[0].rows[0].cells[0].text
            s1_cell_text = section_1_header.tables[0].rows[0].cells[0].text
            
            if s0_cell_text == s1_cell_text and '[TR:' in s0_cell_text:
                print(f"[PASS] {test['name']}")
                print(f"  Section 0 header: {s0_cell_text[:40]}")
                print(f"  Section 1 header: {s1_cell_text[:40]}")
                passed += 1
            else:
                print(f"[FAIL] {test['name']}")
                print(f"  Expected: Same translated text in both headers")
                print(f"  Section 0: {s0_cell_text}")
                print(f"  Section 1: {s1_cell_text}")
                failed += 1
        else:
            print(f"[FAIL] {test['name']}")
            print(f"  Expected: Both headers have tables")
            failed += 1
    
    print(f"\n{passed} passed, {failed} failed\n")
    assert failed == 0, f"{failed} test cases failed"


def test_body_paragraphs_still_translated():
    print("\n=== Testing body paragraphs still translated ===\n")
    
    test_cases = [
        {'name': 'EN fixture - body paragraphs translated', 'fixture': 'test_document_structure_en.docx', 'source_lang': 'en'},
        {'name': 'FR fixture - body paragraphs translated', 'fixture': 'test_document_structure_fr.docx', 'source_lang': 'fr'},
    ]
    
    passed = 0
    failed = 0
    
    for test in test_cases:
        doc, mock = _run_translation(test['fixture'], test['source_lang'])
        
        body_para_indices = [0, 2, 4]
        all_ok = True
        
        for para_idx in body_para_indices:
            if para_idx < len(doc.paragraphs):
                para_text = doc.paragraphs[para_idx].text
                if '[TR:' not in para_text:
                    print(f"[FAIL] {test['name']}")
                    print(f"  Paragraph {para_idx} not translated: {para_text}")
                    all_ok = False
                    break
            else:
                print(f"[FAIL] {test['name']}")
                print(f"  Paragraph {para_idx} does not exist")
                all_ok = False
                break
        
        if all_ok:
            print(f"[PASS] {test['name']}")
            print(f"  Body paragraphs at indices {body_para_indices} all translated")
            passed += 1
        else:
            failed += 1
    
    print(f"\n{passed} passed, {failed} failed\n")
    assert failed == 0, f"{failed} test cases failed"


def test_body_table_still_translated():
    print("\n=== Testing body table still translated ===\n")
    
    test_cases = [
        {'name': 'EN fixture - body table translated', 'fixture': 'test_document_structure_en.docx', 'source_lang': 'en'},
        {'name': 'FR fixture - body table translated', 'fixture': 'test_document_structure_fr.docx', 'source_lang': 'fr'},
    ]
    
    passed = 0
    failed = 0
    
    for test in test_cases:
        doc, mock = _run_translation(test['fixture'], test['source_lang'])
        
        if doc.tables:
            table = doc.tables[0]
            all_cells_ok = True
            
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    if cell_text.strip() and '[TR:' not in cell_text:
                        print(f"[FAIL] {test['name']}")
                        print(f"  Cell not translated: {cell_text}")
                        all_cells_ok = False
                        break
                if not all_cells_ok:
                    break
            
            if all_cells_ok:
                print(f"[PASS] {test['name']}")
                print(f"  Body table cells all translated")
                passed += 1
            else:
                failed += 1
        else:
            print(f"[FAIL] {test['name']}")
            print(f"  Expected: Document has tables")
            failed += 1
    
    print(f"\n{passed} passed, {failed} failed\n")
    assert failed == 0, f"{failed} test cases failed"
