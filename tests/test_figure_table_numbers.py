import os
import tempfile
import pytest
from docx import Document
from scitrans.translate.txt_document import translate_txt_document
from scitrans.translate.word_document import translate_word_document


FIXTURE_DIR = os.path.join(os.path.dirname(__file__), 'fixtures')


class MockTranslator:
    def __init__(self):
        self.call_count = 0
        self.source_texts = []

    def translate_with_best_model(self, text, source_lang, target_lang, use_find_replace, idx, **kwargs):
        self.call_count += 1
        self.source_texts.append(text)
        # Simulate realistic French translation that preserves Figure/Table labels
        translated = text
        translated = translated.replace('Figure 1.', 'Figure 1.')
        translated = translated.replace('Table 1.', 'Tableau 1.')
        # Add a marker so we know translation happened
        translated = f"[TR:{translated}]"
        return {"translated_text": translated}


# ---------------------------------------------------------------------------
# TXT fixture tests (should pass — currently working)
# ---------------------------------------------------------------------------

def test_txt_figure_number_preserved():
    fixture_path = os.path.join(FIXTURE_DIR, 'test_figure_table_numbers.txt')

    temp_output = tempfile.NamedTemporaryFile(suffix='.txt', delete=False)
    temp_output.close()
    output_path = temp_output.name

    mock = MockTranslator()

    translate_txt_document(
        input_text_file=fixture_path,
        output_text_file=output_path,
        source_lang='en',
        use_find_replace=False,
        translation_manager=mock
    )

    with open(output_path, 'r', encoding='utf-8') as f:
        translated_text = f.read()
    os.remove(output_path)

    assert 'Figure 1.' in translated_text, \
        f"Translated text should contain 'Figure 1.', got: {translated_text[:200]}"


def test_txt_table_number_translated():
    fixture_path = os.path.join(FIXTURE_DIR, 'test_figure_table_numbers.txt')

    temp_output = tempfile.NamedTemporaryFile(suffix='.txt', delete=False)
    temp_output.close()
    output_path = temp_output.name

    mock = MockTranslator()

    translate_txt_document(
        input_text_file=fixture_path,
        output_text_file=output_path,
        source_lang='en',
        use_find_replace=False,
        translation_manager=mock
    )

    with open(output_path, 'r', encoding='utf-8') as f:
        translated_text = f.read()
    os.remove(output_path)

    assert 'Tableau 1.' in translated_text, \
        f"Translated text should contain 'Tableau 1.', got: {translated_text[:200]}"


# ---------------------------------------------------------------------------
# DOCX fixture tests (expected to fail until code is updated)
# ---------------------------------------------------------------------------

def _run_docx_translation():
    fixture_path = os.path.join(FIXTURE_DIR, 'test_figure_table_numbers.docx')

    temp_output = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    temp_output.close()
    output_path = temp_output.name

    mock = MockTranslator()

    translate_word_document(
        input_docx_file=fixture_path,
        output_docx_file=output_path,
        source_lang='en',
        use_find_replace=False,
        translation_manager=mock
    )

    doc = Document(output_path)
    os.remove(output_path)

    # Clean up notes file if created
    notes_path = os.path.splitext(output_path)[0] + '_translation_notes.docx'
    if os.path.exists(notes_path):
        os.remove(notes_path)

    return doc, mock


def test_docx_figure_number_preserved():
    doc, _ = _run_docx_translation()

    all_text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'Figure 1.' in all_text, \
        f"Translated docx should contain 'Figure 1.', got: {all_text[:200]}"


def test_docx_table_number_translated():
    doc, _ = _run_docx_translation()

    all_text = '\n'.join(p.text for p in doc.paragraphs)
    assert 'Tableau 1.' in all_text, \
        f"Translated docx should contain 'Tableau 1.', got: {all_text[:200]}"


def test_docx_caption_style_formatting():
    doc, _ = _run_docx_translation()

    # The translated paragraphs should use the "Caption - Figure" style
    # which has default formatting (italic, 10pt)
    caption_paras = [p for p in doc.paragraphs if p.text.strip()]
    assert len(caption_paras) > 0, "Should have at least one non-empty paragraph"

    for para in caption_paras:
        style_name = para.style.name
        assert style_name == 'Caption - Figure', \
            f"Expected 'Caption - Figure' style, got '{style_name}'"
